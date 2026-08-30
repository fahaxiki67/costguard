"""成果导出（Phase 7）：Excel 审核底稿与各类汇总表。

WPS 兼容纪律：
- 只用 openpyxl 基础特性（值/公式/列宽/数字格式），不使用条件 XML 扩展；
- 金额格式 "#,##0.00"；审核底稿保留公式（合价=数量×单价、差异=对比列）；
- 全部导出写入 <project>/exports/，绝不写原始文件（原则 2/3）。
"""
from __future__ import annotations

import json
import re
import sqlite3
from datetime import datetime
from decimal import Decimal
from pathlib import Path

from openpyxl import Workbook
from openpyxl.formatting.rule import CellIsRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from costguard.core.engine.aggregate import aggregate_project, assess_amount, group_key_of
from costguard.core.engine.money import NotANumberError, round2, to_decimal
from costguard.core.engine.settlement_io import pending_sheet_count

D = Decimal

HEADER_FILL = PatternFill("solid", fgColor="DDEBF7")
HEADER_FONT = Font(bold=True)
MONEY_FMT = "#,##0.00"
_THIN_SIDE = Side(style="thin", color="808080")
TABLE_BORDER = Border(
    left=_THIN_SIDE, right=_THIN_SIDE, top=_THIN_SIDE, bottom=_THIN_SIDE
)
CENTER_WRAP = Alignment(horizontal="center", vertical="center", wrap_text=True)
DIRECTION_LABELS = {
    "upward": "对上结算",
    "downward": "对下结算",
    "unknown": "未标记",
}
RULE_ZH_CN = {
    "qty_price_amount_mismatch": "工程量×单价与合价不一致",
    "rounding_difference": "舍入差异",
    "text_number_in_value_col": "数值列存在文本数字",
    "unparsed_number": "数值无法解析",
    "formula_no_cache": "公式缺少缓存值",
    "formula_error": "公式错误",
    "duplicate_item": "重复明细",
    "suspected_duplicate_settlement": "疑似重复结算",
    "price_changed": "单价变化",
    "price_abnormal_change": "单价异常变化",
    "tax_rate_changed": "税率变化",
    "tax_mode_mixed": "计税口径混用",
    "unit_changed": "计量单位变化",
    "qty_sudden_change": "工程量突变",
    "missing_key_fields": "关键字段缺失",
    "summary_mismatch": "汇总金额不一致",
    "summary_missing_data": "汇总数据缺失",
    "negative_quantity": "负数工程量",
    "orphan_numeric_row": "无名称数值行",
    "same_code_diff_name": "同编码不同名称",
    "same_name_diff_code": "同名称不同编码",
    "large_round_amount": "大额整数金额",
    "header_needs_review": "表头识别待复核",
    "missing_columns": "缺少必需列",
    "missing_key_column": "缺少关键列",
}
SUBJECT_ZH = {"line_item": "清单行", "period": "期次", "sheet": "工作表", "project": "项目"}
ANOMALY_STATUS_ZH = {
    "open": "待处理",
    "resolved": "已处理",
    "verified_no_issue": "已核实无问题",
    "supplemented": "已补资料",
    "corrected": "已修正",
    "deferred": "暂不处理",
}


def _business_direction(value: str | None, *, project_level: bool = False) -> str:
    """普通成果列的安全方向文案；绝不把未知内部枚举原样导出。"""
    if value in DIRECTION_LABELS:
        return DIRECTION_LABELS[value]
    return "项目级" if project_level else "未标记"


def _business_rule(value: str | None) -> str:
    return RULE_ZH_CN.get(value or "", "其他审核问题")


def _business_subject(value: str | None) -> str:
    return SUBJECT_ZH.get(value or "", "其他对象")


def _business_status(value: str | None) -> str:
    return ANOMALY_STATUS_ZH.get(value or "", "待人工确认")


def _normalize_business_text(value: str | None) -> str:
    """兼容旧库已生成的短方向词/英文状态摘要。

    新写入内容已经在核心层使用完整业务词；这里的归一化保证旧项目重开后
    重新导出也不会把历史短词带入普通成果列。
    """
    text = str(value or "")
    text = text.replace("[对上]", "[对上结算]").replace("[对下]", "[对下结算]")
    text = text.replace("对上双向校核", "对上结算双向校核")
    text = text.replace("对下双向校核", "对下结算双向校核")
    # 仅修复明确的前缀/短词场景，避免对已经完整的“对上结算”重复替换。
    text = re.sub(r"(?<!结算)对上(?=第|清单|编码|名称|各期)", "对上结算", text)
    text = re.sub(r"(?<!结算)对下(?=第|清单|编码|名称|各期)", "对下结算", text)
    return text


def _style_header(ws, row: int, cols: int):
    for c in range(1, cols + 1):
        cell = ws.cell(row=row, column=c)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(horizontal="center", vertical="center")


def _autowidth(ws):
    for col in ws.columns:
        width = max((len(str(c.value or "")) for c in col), default=8)
        ws.column_dimensions[get_column_letter(col[0].column)].width = min(max(width + 2, 8), 60)


def _style_used_range(ws) -> None:
    """给导出工作表的有效矩形范围统一添加边框、居中和自动换行。"""
    for row in ws.iter_rows(
        min_row=1, max_row=ws.max_row, min_col=1, max_col=ws.max_column
    ):
        for cell in row:
            cell.border = TABLE_BORDER
            cell.alignment = CENTER_WRAP


def _prepare_data_sheet(ws, *, header_row: int = 1) -> None:
    """为可审阅工作表启用冻结表头、筛选、打印重复表头和弱网格显示。"""
    ws.sheet_view.showGridLines = False
    if ws.max_row >= header_row + 1 and ws.max_column:
        ws.freeze_panes = ws.cell(row=header_row + 1, column=1).coordinate
        ws.auto_filter.ref = (
            f"A{header_row}:{get_column_letter(ws.max_column)}{ws.max_row}"
        )
        ws.print_title_rows = f"{header_row}:{header_row}"


def _link_evidence_references(wb: Workbook) -> None:
    """把各成果页 Evidence ID 变为工作簿内跳转链接。"""
    evidence_ws = wb["证据索引"] if "证据索引" in wb.sheetnames else None
    if evidence_ws is None:
        return
    row_by_id = {
        evidence_ws.cell(row=row, column=1).value: row
        for row in range(2, evidence_ws.max_row + 1)
        if evidence_ws.cell(row=row, column=1).value is not None
    }
    for ws in wb.worksheets:
        if ws.title == "证据索引":
            continue
        for row in range(2, ws.max_row + 1):
            for col in range(1, ws.max_column + 1):
                cell = ws.cell(row=row, column=col)
                value = cell.value
                header = str(ws.cell(row=1, column=col).value or "")
                evidence_column = "证据" in header or "出处" in header
                if not evidence_column and not (
                    isinstance(value, str) and "Evidence ID" in value
                ):
                    continue
                target_id = value if isinstance(value, int) else None
                if target_id is None and isinstance(value, str):
                    match = re.search(r"Evidence ID\s+(\d+)", value)
                    target_id = int(match.group(1)) if match else None
                target_row = row_by_id.get(target_id)
                if target_row is not None:
                    cell.hyperlink = f"#'证据索引'!A{target_row}"


# openpyxl 原生支持 Decimal（XML 层写精确十进制字符串），当前序列化无需 float。
# 若未来更换写入引擎导致必须经 float，本常量置 True：金额先 round2 再转 float，
# 且转换只允许发生在 _num 这一个边界（监督门槛 4）。
_SERIALIZE_NEEDS_FLOAT = False


def _num(value, money: bool = False):
    """序列化边界：业务值 → xlsx 数值。全模块唯一的"值 → 单元格"转换点。

    - Decimal 直写，XML 层为精确十进制字符串，无二进制误差，业务值保真
      （原表合价等证据值不得因序列化改变）；
    - money=True 仅为金额列语义标记：若 _SERIALIZE_NEEDS_FLOAT（必须转 float 的
      写入引擎），则先 round2 再转 float；
    - 缺失/不可解析一律 None（待补资料），绝不补 0；
    - 禁止在业务计算层做任何 float 转换。
    """
    if value is None:
        return None
    if isinstance(value, Decimal):
        d = value
    elif isinstance(value, str):
        try:
            d = to_decimal(value)
        except NotANumberError:
            return None
    else:
        d = Decimal(repr(value))
    if _SERIALIZE_NEEDS_FLOAT:
        return float(round2(d)) if money else float(d)
    return d


def _fetch_periods(conn, project_id):
    return conn.execute(
        "SELECT id, period_no, title, direction, contract_party, tax_mode FROM settlement_periods"
        " WHERE project_id=? ORDER BY CASE direction WHEN 'upward' THEN 0 "
        "WHEN 'downward' THEN 1 ELSE 2 END, period_no, id",
        (project_id,),
    ).fetchall()


def _project_directions(conn: sqlite3.Connection, project_id: int) -> list[str]:
    """返回项目实际存在的方向；空项目仍给出未标记方向占位。"""
    rows = conn.execute(
        """SELECT DISTINCT COALESCE(direction, 'unknown') AS direction
           FROM settlement_periods WHERE project_id=?
           ORDER BY CASE COALESCE(direction, 'unknown')
             WHEN 'upward' THEN 0 WHEN 'downward' THEN 1 ELSE 2 END""",
        (project_id,),
    ).fetchall()
    return [r["direction"] for r in rows] or ["unknown"]


def export_settlement_summary(conn: sqlite3.Connection, project_id: int, wb: Workbook,
                              direction: str | None = None) -> str:
    """结算累计表（对上/对下各自独立累计，绝不混入另一方向数据）。"""
    if direction is None:
        raise ValueError("结算累计表必须明确 direction，禁止生成跨方向混合汇总")
    periods = [
        p for p in _fetch_periods(conn, project_id)
        if (p["direction"] or "unknown") == direction
    ]
    aggs = aggregate_project(conn, project_id, direction=direction)
    ws = wb.create_sheet(f"{_business_direction(direction)}累计表")
    header = ["清单编码", "清单名称", "单位"] + [f"第{p['period_no']}期金额" for p in periods] + \
             ["累计数量", "累计金额", "加权平均单价", "状态"]
    ws.append(header)
    _style_header(ws, 1, len(header))
    for agg in aggs:
        row = [agg.code, agg.name, ""]
        for p in periods:
            pp = agg.per_period.get(p["id"])  # period_id 键：防对上/对下同期号串表
            row.append(_num(pp["effective_amount"], money=True) if pp else None)
        row.append(_num(agg.cum_qty))
        row.append(_num(agg.cum_amount, money=True))
        row.append(_num(agg.wavg_price, money=True))
        status = {"ok": "正常", "incomplete": "待补资料", "incomparable": "不可比"}[agg.status]
        row.append(status + ("；".join(agg.warnings[:2]) if agg.warnings else ""))
        ws.append(row)
    _autowidth(ws)
    n = ws.max_row
    if n > 1:
        for r in range(2, n + 1):
            for c in range(4, len(header)):
                ws.cell(row=r, column=c).number_format = MONEY_FMT
    return ws.title


def _diff_series(conn: sqlite3.Connection, project_id: int, field: str) -> list[dict]:
    """差异表数据序列：按 (direction, period_id, item_key) 取可追溯原始值。

    - 数量/金额：同期同组多行求和（Decimal）；单位不一致时该期标记不可比；
    - 单价：同期同组若存在多个不同价 → 标记"多价待复核"，绝不平均；
    - 缺失值不补 0：整期无有效值 → None（待补资料）。
    """
    rows = conn.execute(
        """SELECT li.code, li.name, li.unit, li.quantity, li.unit_price, li.amount,
                  li.flags_json, sp.period_no AS pno, sp.direction AS dir
           FROM line_items li JOIN settlement_periods sp ON sp.id = li.period_id
           WHERE sp.project_id=? ORDER BY sp.period_no, li.id""",
        (project_id,),
    ).fetchall()
    # 归组键与 aggregate.group_key_of 同口径：code 优先；同方向同码跨期改名
    # 仍是同一序列（名称变化逐期保留并加变更提示，不按名称拆分）。
    series: dict[tuple[str, str], dict[int, dict]] = {}
    series_names: dict[tuple[str, str], set] = {}
    for r in rows:
        flags = json.loads(r["flags_json"] or "{}")
        if flags.get("subtotal"):
            continue
        key = (r["dir"] or "unknown", "code:" + r["code"] if r["code"] else "name:" + (r["name"] or ""))
        by_period = series.setdefault(key, {})
        series_names.setdefault(key, set()).add(r["name"] or "")
        pp = by_period.setdefault(int(r["pno"]), {
            "qty": None, "amount": None, "effective_amount": None,
            "amount_source": "missing", "prices": set(), "units": set(), "qty_missing": False,
            "period_no": int(r["pno"]), "name": r["name"] or "",
        })
        if r["quantity"] is not None:
            try:
                q = D(r["quantity"])
            except Exception:
                q = None
            if q is not None:
                pp["qty"] = q if pp["qty"] is None else pp["qty"] + q
        else:
            pp["qty_missing"] = True
        assessment, _qty_missing, _price_missing = assess_amount(r)
        if assessment.raw is not None:
            a = assessment.raw
            pp["amount"] = a if pp["amount"] is None else pp["amount"] + a
        if assessment.effective is not None:
            a = assessment.effective
            pp["effective_amount"] = (
                a if pp["effective_amount"] is None else pp["effective_amount"] + a
            )
            if pp["amount_source"] in ("missing", assessment.source):
                pp["amount_source"] = assessment.source
            else:
                pp["amount_source"] = "mixed"
        if r["unit_price"]:
            try:
                pp["prices"].add(D(r["unit_price"]))
            except Exception:
                pass
        if r["unit"]:
            pp["units"].add(_norm_unit_local(r["unit"]))
    out = []
    for key, by_period in series.items():
        direction, key_str = key
        code = key_str[5:] if key_str.startswith("code:") else ""
        out.append({
            "direction": direction, "code": code,
            "names": series_names[key], "by_period": by_period,
        })
    return sorted(out, key=lambda x: (x["direction"], x["code"]))


def _norm_unit_local(u: str | None) -> str:
    if not u:
        return ""
    u = u.strip()
    for canonical, aliases in UNIT_ALIAS_EXPORT.items():
        if u == canonical or u in aliases:
            return canonical
    return u


UNIT_ALIAS_EXPORT = {
    "m3": {"m³", "立方米", "立米", "M3", "m^3"},
    "m2": {"m²", "平方米", "平米", "M2", "m^2"},
}


def export_diff_sheets(conn: sqlite3.Connection, project_id: int, wb: Workbook) -> None:
    """单价/工程量/金额 差异表：方向隔离的跨期比较。

    - 按 direction 分组，同方向内按 period_no 比较；对上与对下绝不互比；
    - 归组 code 优先（同码跨期改名仍是同一序列，名称逐期保留并加变更提示）；
    - 单价来自可追溯原始 unit_price；同期同组多价 → 标"多价待复核"，不得平均；
    - 不可比/缺失期是断点：清空比较基准，后续可比期作为新首期，绝不跨越强行比较；
    - 差异列保留公式；缺失不补 0。
    """
    titles = (("单价差异表", "unit_price"), ("工程量差异表", "quantity"), ("金额差异表", "amount"))
    dir_zh = DIRECTION_LABELS
    all_series = {f: _diff_series(conn, project_id, f) for _t, f in titles}

    for title, field in titles:
        ws = wb.create_sheet(title)
        ws.append(["方向", "清单编码", "清单名称", "期间", "本期值", "上期值", "差异", "差异率"])
        _style_header(ws, 1, 8)
        series = all_series[field]
        r = 1
        for item in series:
            ordered = sorted(item["by_period"].items(), key=lambda kv: kv[1]["period_no"])
            names = item["names"]
            first_name = item["by_period"][ordered[0][0]].get("name", "") if ordered else ""
            prev: Decimal | None = None  # 上一可比期数值；不可比/缺失/单位变化即断点
            prev_units: set | None = None
            for _pid, pp in ordered:
                r += 1
                pno = pp["period_no"]
                cur_name = pp.get("name", "")
                # 名称变更提示：非首现名称的期标注（首期/首现名保持原样，变更可追溯）
                if len(names) > 1 and cur_name != first_name:
                    name_cell = f"{cur_name}（名称变更）"
                else:
                    name_cell = cur_name
                if field == "unit_price":
                    if len(pp["prices"]) > 1:
                        cur_val: Decimal | str | None = "多价待复核（不可比，不平均）"
                    else:
                        cur_val = _num(next(iter(pp["prices"])), money=True) if pp["prices"] else None
                elif field == "quantity":
                    if len(pp["units"]) > 1:
                        cur_val = "单位不一致（不可比）"
                    else:
                        cur_val = _num(pp["qty"]) if pp["qty"] is not None else None
                else:
                    cur_val = (
                        _num(pp["effective_amount"], money=True)
                        if pp["effective_amount"] is not None else None
                    )

                ws.cell(row=r, column=1, value=dir_zh.get(item["direction"], item["direction"]))
                ws.cell(row=r, column=2, value=item["code"])
                ws.cell(row=r, column=3, value=name_cell)
                ws.cell(row=r, column=4, value=f"第{pno}期")
                ws.cell(row=r, column=5, value=cur_val)

                # 跨期单位变化：与上一可比期单位无交集 → 不可比断点
                unit_mismatch = (
                    isinstance(cur_val, Decimal)
                    and prev_units is not None
                    and pp["units"] is not None
                    and prev_units
                    and not (pp["units"] & prev_units)
                )
                if isinstance(cur_val, Decimal) and unit_mismatch:
                    ws.cell(row=r, column=7, value="不可比（与上期单位不一致）")
                    ws.cell(row=r, column=8, value="不可比")
                    prev = None
                    prev_units = set(pp["units"])
                elif isinstance(cur_val, Decimal):
                    if prev is not None:
                        ws.cell(row=r, column=6, value=prev)
                        ws.cell(row=r, column=7, value=f"=E{r}-F{r}")
                        if prev != 0:
                            ws.cell(row=r, column=8, value=f'=IF(F{r}=0,"不可比",(E{r}-F{r})/F{r})')
                            ws.cell(row=r, column=8).number_format = "0.00%"
                        else:
                            ws.cell(row=r, column=8, value="不可比（上期为 0）")
                    else:
                        ws.cell(row=r, column=8, value="首期（无上期可比）")
                    prev = cur_val
                    prev_units = set(pp["units"]) if pp["units"] else None
                else:
                    # 不可比/缺失期：断点——清空基准，绝不与前后期强行比较
                    ws.cell(row=r, column=7, value="不可比" if isinstance(cur_val, str) else "待补资料")
                    ws.cell(row=r, column=8, value="不可比")
                    prev = None
                    prev_units = None
                for c in (5, 6, 7):
                    if isinstance(ws.cell(row=r, column=c).value, Decimal):
                        ws.cell(row=r, column=c).number_format = MONEY_FMT
        _autowidth(ws)


def _aggregate_by_direction(conn: sqlite3.Connection, project_id: int, direction: str) -> dict[str, dict]:
    """按方向聚合：item_key -> {qty, amount, names}。缺失值不补 0。"""
    rows = conn.execute(
        """SELECT li.id, li.code, li.name, li.unit, li.quantity, li.unit_price, li.amount,
                  li.flags_json FROM line_items li
           JOIN settlement_periods sp ON sp.id = li.period_id
           WHERE sp.project_id=? AND sp.direction=?""",
        (project_id, direction),
    ).fetchall()
    out: dict[str, dict] = {}
    for r in rows:
        flags = json.loads(r["flags_json"] or "{}")
        if flags.get("subtotal"):
            continue
        key = group_key_of(r["code"], r["name"] or "")
        agg = out.setdefault(key, {"qty": None, "amount": None, "names": set()})
        agg["names"].add(r["name"] or "")
        assessment, _qty_missing, _price_missing = assess_amount(r)
        for field, val in (("qty", r["quantity"]), ("amount", assessment.effective)):
            if val is not None:  # "0" 是有效值参与累计；缺失(None)不参与也不补 0
                try:
                    d = D(val)
                except Exception:
                    continue
                agg[field] = d if agg[field] is None else agg[field] + d
    return out


def export_updown_comparison(conn: sqlite3.Connection, project_id: int, wb: Workbook) -> None:
    """对上对下对比表：同一清单的对上累计 vs 对下累计（差异列为公式）。"""
    ws = wb.create_sheet("对上对下对比表")
    ws.append(["清单编码", "清单名称", "对上累计数量", "对下累计数量", "对上累计金额",
               "对下累计金额", "金额差异(公式)", "口径说明"])
    _style_header(ws, 1, 8)
    up = _aggregate_by_direction(conn, project_id, "upward")
    down = _aggregate_by_direction(conn, project_id, "downward")
    if not up and not down:
        ws.append(["", "无可比数据：请先在期次管理中标记对上/对下方向", None, None, None, None, None,
                   "期次方向未标记时无法对比（不可比，不做强行比较）"])
    keys = sorted(set(up) | set(down))
    r = 1
    for key in keys:
        r += 1
        u, d = up.get(key), down.get(key)
        code = key[5:] if key.startswith("code:") else ""
        names = sorted((u or d or {}).get("names") or {""})
        ws.cell(row=r, column=1, value=code)
        ws.cell(row=r, column=2, value=names[0])
        if u and u["qty"] is not None:
            ws.cell(row=r, column=3, value=_num(u["qty"]))
        if d and d["qty"] is not None:
            ws.cell(row=r, column=4, value=_num(d["qty"]))
        if u and u["amount"] is not None:
            ws.cell(row=r, column=5, value=_num(u["amount"], money=True))
        if d and d["amount"] is not None:
            ws.cell(row=r, column=6, value=_num(d["amount"], money=True))
        # 差异公式：两侧都有值才算（一侧缺失标记待补资料，不补 0）
        if u and d and u["amount"] is not None and d["amount"] is not None:
            ws.cell(row=r, column=7, value=f"=E{r}-F{r}")
            ws.cell(row=r, column=7).number_format = MONEY_FMT
        else:
            ws.cell(row=r, column=8, value="待补资料（一侧缺失，不做比较）")
        if u and d and set(map(str, u["names"])) != set(map(str, d["names"])):
            ws.cell(row=r, column=8, value="两侧名称不一致，请核实归组")
        for c in (3, 4, 5, 6):
            ws.cell(row=r, column=c).number_format = MONEY_FMT
    _autowidth(ws)


def export_anomaly_lists(conn: sqlite3.Connection, project_id: int, wb: Workbook) -> None:
    """异常清单 + 待核实事项清单。"""
    anomaly_rows = conn.execute(
        """SELECT a.id, a.rule_id, a.severity, a.subject_type, a.subject_id,
                  a.evidence_id, a.message, a.status,
                  COALESCE(sp_item.direction, sp_period.direction, sp_sheet.direction, '')
                    AS direction
           FROM anomalies a
           LEFT JOIN line_items li
             ON a.subject_type='line_item' AND li.id=a.subject_id
           LEFT JOIN settlement_periods sp_item ON sp_item.id=li.period_id
           LEFT JOIN settlement_periods sp_period
             ON a.subject_type='period' AND sp_period.id=a.subject_id
           LEFT JOIN raw_sheets rs
             ON a.subject_type='sheet' AND rs.id=a.subject_id
           LEFT JOIN settlement_periods sp_sheet ON sp_sheet.id=rs.period_id
           WHERE a.project_id=? ORDER BY CASE a.severity
           WHEN 'high' THEN 0 WHEN 'medium' THEN 1 WHEN 'low' THEN 2 ELSE 3 END, a.id""",
        (project_id,),
    ).fetchall()
    sev_zh = {"high": "高", "medium": "中", "low": "低", "info": "提示"}
    ws = wb.create_sheet("异常清单")
    ws.append(["编号", "方向", "规则", "级别", "对象", "说明", "证据ID", "状态", "规则代码"])
    _style_header(ws, 1, 9)
    for r in anomaly_rows:
        sev = sev_zh.get(r["severity"], "其他")
        direction = _business_direction(r["direction"], project_level=not r["direction"])
        subject = _business_subject(r["subject_type"])
        status = _business_status(r["status"])
        ws.append([r["id"], direction, _business_rule(r["rule_id"]), sev,
                   f"{subject}#{r['subject_id']}",
                   _normalize_business_text(r["message"]), r["evidence_id"], status, r["rule_id"]])
    _autowidth(ws)
    # 原始规则编码仅供高级排查，默认隐藏，避免普通业务界面直接暴露开发字段；
    # 列仍保留以便技术人员在需要时取消隐藏并追溯。
    ws.column_dimensions["I"].hidden = True

    ws2 = wb.create_sheet("待核实事项清单")
    ws2.append(["编号", "方向", "类别", "说明", "证据ID", "规则代码"])
    _style_header(ws2, 1, 6)
    idx = 1
    for r in anomaly_rows:
        if r["severity"] not in {"high", "medium"} or r["status"] not in {"open", "deferred"}:
            continue
        direction = _business_direction(r["direction"], project_level=not r["direction"])
        message = _normalize_business_text(r["message"])
        if r["status"] == "deferred":
            message = f"【暂不处理】{message}"
        ws2.append([idx, direction, _business_rule(r["rule_id"]),
                    message, r["evidence_id"], r["rule_id"]])
        idx += 1
    for direction in _project_directions(conn, project_id):
        for agg in aggregate_project(conn, project_id, direction=direction):
            if agg.status in ("incomplete", "incomparable"):
                for w in agg.warnings:
                    ws2.append([
                        idx,
                        _business_direction(direction),
                        "汇总校核",
                        _normalize_business_text(f"「{agg.name}」{w}"),
                        None,
                    ])
                    idx += 1
    _autowidth(ws2)
    ws2.column_dimensions["F"].hidden = True


def export_contract_risks(conn: sqlite3.Connection, project_id: int, wb: Workbook) -> None:
    ws = wb.create_sheet("合同风险清单")
    ws.append(["编号", "级别", "合同文档", "风险说明", "证据ID"])
    _style_header(ws, 1, 5)
    rows = conn.execute(
        """SELECT a.id, a.severity, a.message, a.evidence_id, cd.title
           FROM anomalies a JOIN contract_docs cd ON cd.id = a.subject_id
           WHERE a.project_id=? AND a.rule_id='contract_risk'""",
        (project_id,),
    ).fetchall()
    for i, r in enumerate(rows, start=1):
        sev = {"high": "高", "medium": "中", "low": "低"}.get(r["severity"], r["severity"])
        ws.append([i, sev, r["title"], r["message"], r["evidence_id"]])
    _autowidth(ws)


def export_evidence_index(conn: sqlite3.Connection, project_id: int, wb: Workbook) -> None:
    ws = wb.create_sheet("证据索引")
    ws.append(["证据ID", "类型", "摘要", "计算过程(JSON)", "来源(JSON)", "生成时间"])
    _style_header(ws, 1, 6)
    for r in conn.execute(
        "SELECT id, kind, summary, steps_json, sources_json, created_at FROM evidence WHERE project_id=? ORDER BY id",
        (project_id,),
    ):
        ws.append([r["id"], r["kind"], _normalize_business_text(r["summary"]),
                   r["steps_json"], r["sources_json"], r["created_at"]])
    _autowidth(ws)


def export_audit_worksheet(conn: sqlite3.Connection, project_id: int, wb: Workbook) -> None:
    """Excel 审核底稿：逐行明细 + 出处 + 双值列（公式 + 程序 Decimal 值）。

    双值纪律（P0-3）：合价列保留公式（Office 打开重算=复核值）；
    程序计算合价列写入引擎 Decimal 精确值——主结论不依赖 Office 重算。
    """
    ws = wb.create_sheet("审核底稿")
    ws.append(["期次", "清单编码", "清单名称", "单位", "数量", "单价", "合价(底稿公式)",
               "程序计算合价", "原表合价", "差异(公式)", "数量出处", "单价出处", "合价出处", "行ID", "方向"])
    _style_header(ws, 1, 15)
    rows = conn.execute(
        """SELECT li.*, sp.period_no AS pno, COALESCE(sp.direction, 'unknown') AS direction
           FROM line_items li
           JOIN settlement_periods sp ON sp.id = li.period_id
           WHERE sp.project_id=?
           ORDER BY CASE COALESCE(sp.direction, 'unknown')
             WHEN 'upward' THEN 0 WHEN 'downward' THEN 1 ELSE 2 END,
             sp.period_no, li.id""",
        (project_id,),
    ).fetchall()
    r = 1
    for row in rows:
        flags = json.loads(row["flags_json"] or "{}")
        if flags.get("subtotal"):
            continue
        source_evidence_id = flags.get("source_evidence_id")
        r += 1
        # 缺失仅按 is None 判断：Decimal("0") 是有效值必须保留（监督门槛 1）
        qty = _num(row["quantity"]) if row["quantity"] is not None else None
        price = _num(row["unit_price"], money=True) if row["unit_price"] is not None else None
        amount = _num(row["amount"], money=True) if row["amount"] is not None else None
        ws.cell(row=r, column=1, value=row["pno"])
        ws.cell(row=r, column=2, value=row["code"])
        ws.cell(row=r, column=3, value=row["name"])
        ws.cell(row=r, column=4, value=row["unit"])
        ws.cell(row=r, column=5, value=qty)
        ws.cell(row=r, column=6, value=price)
        program_amount = None
        if qty is not None and price is not None:
            ws.cell(row=r, column=7, value=f"=E{r}*F{r}")  # 保留公式（Excel 复核值）
            from costguard.core.engine.money import money_mul

            program_amount = money_mul(qty, price)  # 引擎 Decimal 精确值
            ws.cell(row=r, column=8, value=program_amount)
        else:
            ws.cell(row=r, column=7, value="待补资料")
            ws.cell(row=r, column=8, value="待补资料")
        ws.cell(row=r, column=9, value=amount)
        if qty is not None and price is not None and amount is not None:
            ws.cell(row=r, column=10, value=f"=ROUND(G{r}-I{r},2)")
        else:
            ws.cell(row=r, column=10, value="不可比")
        ws.cell(row=r, column=14, value=row["id"])  # 行ID：唯一标识，复核回溯用
        ws.cell(
            row=r,
            column=15,
            value=_business_direction(row["direction"]),
        )
        for col, evid in ((11, "qty_evid"), (12, "price_evid"), (13, "amount_evid")):
            ev = row[evid]
            if ev:
                # 当前导入器保存字段来源 JSON；兼容历史库中可能存在的整数
                # evidence_id 或损坏的旧值，单个来源异常不得让整本底稿导出失败。
                try:
                    e = json.loads(ev) if isinstance(ev, str) else None
                except (TypeError, json.JSONDecodeError):
                    e = None
                if isinstance(e, dict):
                    row_no, col_no = e.get("row", "—"), e.get("col", "—")
                    raw = str(e.get("raw", e.get("value", "—")))[:30]
                    evidence_text = (
                        f"Evidence ID {e.get('evidence_id') or source_evidence_id} · "
                        if (e.get("evidence_id") or source_evidence_id) else ""
                    )
                    ws.cell(row=r, column=col, value=f"{evidence_text}行{row_no}列{col_no}: {raw}")
                else:
                    ws.cell(row=r, column=col, value=f"Evidence ID {ev}")
        for c in (5, 6, 7, 8, 9, 10):
            ws.cell(row=r, column=c).number_format = MONEY_FMT
    _autowidth(ws)


def _tax_mode_status(conn: sqlite3.Connection, project_id: int) -> str:
    """各期单价含税口径状态（供摘要声明范围）。"""
    modes: dict[str, set] = {}
    for h in conn.execute(
        """SELECT th.col_map_json, th.header_row_lo, th.header_row_hi, rs.id AS sid
           FROM table_headers th JOIN raw_sheets rs ON rs.id = th.sheet_id
           JOIN settlement_periods sp ON sp.id = rs.period_id WHERE sp.project_id=?""",
        (project_id,),
    ):
        col_map = json.loads(h["col_map_json"])
        if not ({"unit_price", "amount"} & set(col_map)):
            continue
        texts = [r["raw_value"] or "" for r in conn.execute(
            "SELECT raw_value FROM raw_cells WHERE sheet_id=? AND row BETWEEN ? AND ?",
            (h["sid"], h["header_row_lo"], h["header_row_hi"]),
        )]
        joined = "".join(texts)
        if "不含税" in joined:
            modes.setdefault("excl_tax", set()).add(h["sid"])
        elif "含税" in joined:
            modes.setdefault("incl_tax", set()).add(h["sid"])
    if not modes:
        return "未识别（表头未标注含税/不含税，请人工确认）"
    if len(modes) > 1:
        return "混用（存在含税与不含税口径，详见异常清单 tax_mode_mixed）"
    return "含税金额/单价" if "incl_tax" in modes else "不含税金额/单价"


def _review_gate_counts(conn: sqlite3.Connection, project_id: int) -> dict[str, int]:
    """集中计算成果页使用的审核完成度闸门，避免不同成果口径漂移。"""
    source_files = conn.execute(
        "SELECT COUNT(*) AS c FROM source_files WHERE project_id=?", (project_id,)
    ).fetchone()["c"]
    pending_sheets = pending_sheet_count(conn, project_id)
    high_unresolved = conn.execute(
        """SELECT COUNT(*) AS c FROM anomalies
           WHERE project_id=? AND severity='high' AND status IN ('open', 'deferred')""",
        (project_id,),
    ).fetchone()["c"]
    open_findings = conn.execute(
        """SELECT COUNT(*) AS c FROM anomalies
           WHERE project_id=? AND status IN ('open', 'deferred')""",
        (project_id,),
    ).fetchone()["c"]
    deferred = conn.execute(
        "SELECT COUNT(*) AS c FROM anomalies WHERE project_id=? AND status='deferred'",
        (project_id,),
    ).fetchone()["c"]
    pending_matches = conn.execute(
        "SELECT COUNT(*) AS c FROM matches WHERE project_id=? AND status='pending'",
        (project_id,),
    ).fetchone()["c"]
    period_count = conn.execute(
        "SELECT COUNT(*) AS c FROM settlement_periods WHERE project_id=?", (project_id,)
    ).fetchone()["c"]
    checked_count = conn.execute(
        "SELECT COUNT(*) AS c FROM crosscheck_results WHERE project_id=?", (project_id,)
    ).fetchone()["c"]
    insufficient = conn.execute(
        """SELECT COUNT(*) AS c FROM crosscheck_results
           WHERE project_id=? AND verification_level='insufficient'""",
        (project_id,),
    ).fetchone()["c"]
    findings = conn.execute(
        """SELECT COUNT(*) AS c FROM crosscheck_results
           WHERE project_id=? AND verification_level='findings'""",
        (project_id,),
    ).fetchone()["c"]
    range_unproven = conn.execute(
        """SELECT COALESCE(SUM(range_unproven_sheets), 0) AS c
           FROM crosscheck_results WHERE project_id=?""",
        (project_id,),
    ).fetchone()["c"]
    return {
        "source_files": int(source_files or 0),
        "pending_sheets": int(pending_sheets or 0),
        "high_unresolved": int(high_unresolved or 0),
        "open_findings": int(open_findings or 0),
        "deferred": int(deferred or 0),
        "pending_matches": int(pending_matches or 0),
        "period_count": int(period_count or 0),
        "checked_count": int(checked_count or 0),
        "insufficient": int(insufficient or 0),
        "findings": int(findings or 0),
        "range_unproven": int(range_unproven or 0),
    }


def export_cover_page(conn: sqlite3.Connection, project_id: int, wb: Workbook) -> None:
    """生成成果首页，首次打开即可看到范围、状态、限制和使用顺序。"""
    ws = wb.create_sheet("封面与说明")
    project = conn.execute(
        "SELECT name FROM projects WHERE id=?", (project_id,)
    ).fetchone()
    project_name = project["name"] if project else "未命名项目"
    gates = _review_gate_counts(conn, project_id)
    unchecked = max(0, gates["period_count"] - gates["checked_count"])
    if not gates["source_files"]:
        status = "尚未导入资料"
    elif not gates["period_count"]:
        status = "暂无可审核结算期次"
    elif any((
        gates["pending_sheets"], gates["high_unresolved"], gates["open_findings"],
        gates["pending_matches"], gates["insufficient"], gates["findings"],
        gates["range_unproven"], unchecked,
    )):
        status = "审核尚未完成"
    else:
        status = "当前未发现主要待处理事项"
    ws.append(["CostGuard Excel 审核底稿"])
    ws.cell(row=1, column=1).font = Font(bold=True, size=18)
    ws.append(["项目名称", project_name])
    ws.append(["生成时间", datetime.now().strftime("%Y-%m-%d %H:%M")])
    ws.append(["成果版本", "v0.1.7 预览候选"])
    ws.append(["审核状态", status])
    completion = (
        f"待确认工作表 {gates['pending_sheets']} 张；高风险未处理 {gates['high_unresolved']} 项；"
        f"待确认匹配 {gates['pending_matches']} 组；尚未校核 {unchecked} 期；"
        f"校核不充分 {gates['insufficient']} 期；取数范围未证明 {gates['range_unproven']} 张工作表"
    )
    if not gates["source_files"]:
        completion = "尚未导入资料；" + completion
    elif not gates["period_count"]:
        completion = "暂无可审核结算期次；" + completion
    ws.append(["审核完成度", completion])
    ws.append([])
    ws.append(["建议使用顺序"])
    ws.append(["1", "先看本页状态和限制，再看《管理层摘要》"])
    ws.append(["2", "在《异常清单》《待核实事项清单》中处理问题并记录原因"])
    ws.append(["3", "在《审核底稿》中核对程序 Decimal 值、Excel 复核公式值和原始合价"])
    ws.append(["4", "按《证据索引》中的 Evidence ID 回溯文件、Sheet、行列和原始值"])
    ws.append([])
    ws.append(["重要说明"])
    ws.append(["缺失数据保持“待补资料”，不可比数据保持“不可比”，不以 0 替代。"])
    ws.append(["“校核充分”仅表示证据条件充分且验证通过，不等同于业务审批或最终结算确认。"])
    ws.append(["程序计算值是主审核依据；公式列供 WPS/Excel 打开后复核，不依赖 Office 首次重算。"])
    ws.append(["WPS、macOS Excel、Windows Excel 真机验证及大规模性能基准仍是发布门槛。"])
    _autowidth(ws)


def export_management_summary(conn: sqlite3.Connection, project_id: int, wb: Workbook) -> None:
    """管理层摘要 sheet：范围、期次与方向、税口径、异常与证据计数。"""
    ws = wb.create_sheet("管理层摘要")
    ws.append(["CostGuard 管理层摘要"])
    ws.cell(row=1, column=1).font = Font(bold=True, size=14)
    direction_stats: dict[str, dict] = {}
    for direction in _project_directions(conn, project_id):
        aggs = aggregate_project(conn, project_id, direction=direction)
        direction_stats[direction] = {
            "groups": len(aggs),
            "total": sum((a.cum_amount for a in aggs if a.cum_amount is not None), D(0)),
            "ok": sum(1 for a in aggs if a.status == "ok"),
            "incomplete": sum(1 for a in aggs if a.status == "incomplete"),
            "incomparable": sum(1 for a in aggs if a.status == "incomparable"),
        }
    sev = {"high": 0, "medium": 0, "low": 0}
    for r in conn.execute("SELECT severity, COUNT(*) c FROM anomalies WHERE project_id=? GROUP BY severity", (project_id,)):
        sev[r["severity"]] = r["c"]
    periods = _fetch_periods(conn, project_id)
    n_up = sum(1 for p in periods if p["direction"] == "upward")
    n_down = sum(1 for p in periods if p["direction"] == "downward")
    n_none = len(periods) - n_up - n_down
    n_ev = conn.execute("SELECT COUNT(*) c FROM evidence WHERE project_id=?", (project_id,)).fetchone()["c"]
    n_src = conn.execute("SELECT COUNT(*) c FROM source_files WHERE project_id=?", (project_id,)).fetchone()["c"]
    data = [
        ("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("—— 统计范围 ——", ""),
        ("已导入原始文件数", n_src),
        ("期次数", len(periods)),
        ("其中：对上结算", n_up),
        ("其中：对下结算", n_down),
        ("其中：未标记", n_none),
        ("—— 金额与状态 ——", ""),
    ]
    for direction, stats in direction_stats.items():
        label = _business_direction(direction)
        data.extend([
            (f"{label}清单组数", stats["groups"]),
            (f"{label}累计金额（可用部分）", _num(round2(stats["total"]))),
            (f"{label}正常清单组", stats["ok"]),
            (f"{label}待补资料清单组", stats["incomplete"]),
            (f"{label}不可比清单组", stats["incomparable"]),
        ])
    data.extend([
        ("—— 口径与风险 ——", ""),
        ("金额/单价税口径", _tax_mode_status(conn, project_id)),
        ("高风险异常数", sev["high"]),
        ("中风险异常数", sev["medium"]),
        ("低风险异常数", sev["low"]),
        ("—— 追溯 ——", ""),
        ("证据记录数", n_ev),
        ("证据索引位置", "本工作簿《证据索引》工作表（evidence ID）"),
    ])
    for k, v in data:
        ws.append([k, v])
    ws.append([])
    ws.append(["说明：本摘要由 CostGuard 自动生成，仅反映已导入数据。缺失数据未补 0；"
               "不可比数据未强行比较；对上、对下和未标记方向分别列示，未作跨方向净额或合计。"])
    ws.append(["自动计算结果必须经过人工复核和业务审批；未经批准，不构成任何真实业务结论，"
               "包括最终结算、责任认定或正式管理结论。"])
    _autowidth(ws)


def export_workbook(conn: sqlite3.Connection, project_id: int, out_dir: Path) -> Path:
    """导出全部报表到一个 xlsx。返回文件路径。"""
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    wb.remove(wb.active)
    export_cover_page(conn, project_id, wb)
    export_management_summary(conn, project_id, wb)
    for direction in _project_directions(conn, project_id):
        export_settlement_summary(conn, project_id, wb, direction=direction)
    export_updown_comparison(conn, project_id, wb)
    export_diff_sheets(conn, project_id, wb)
    export_anomaly_lists(conn, project_id, wb)
    export_contract_risks(conn, project_id, wb)
    export_evidence_index(conn, project_id, wb)
    export_audit_worksheet(conn, project_id, wb)
    for ws in wb.worksheets:
        _style_used_range(ws)
        if ws.title not in {"封面与说明", "管理层摘要"}:
            _prepare_data_sheet(ws)
    audit = wb["审核底稿"]
    if audit.max_row >= 2:
        # 差异非零时在 WPS/Excel 中以浅红色提示；缺失/不可比文字仍保持原状态。
        audit.conditional_formatting.add(
            f"J2:J{audit.max_row}",
            CellIsRule(operator="notEqual", formula=["0"],
                       fill=PatternFill("solid", fgColor="FDECEC")),
        )
    _link_evidence_references(wb)
    # 不能在当前运行环境伪造 Excel 公式缓存值；明确要求 Office 首次打开时
    # 全量重算，并把程序 Decimal 列作为主审核依据，降低 WPS/Excel 显示旧缓存
    # 或 0.00 的风险。WPS、macOS Excel、Windows Excel 真机复核仍是发布门槛。
    wb.calculation.calcMode = "auto"
    wb.calculation.fullCalcOnLoad = True
    wb.calculation.forceFullCalc = True
    wb.calculation.calcOnSave = True
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = out_dir / f"CostGuard审核底稿_{stamp}.xlsx"
    wb.save(path)
    return path


def export_management_summary_docx(conn: sqlite3.Connection, project_id: int, out_dir: Path) -> Path:
    """管理层摘要 Word 版：首屏状态、关键指标、Top 风险和限制均可追溯。"""
    import docx as docx_lib
    from docx.oxml.ns import qn
    from docx.shared import Pt

    project = conn.execute("SELECT name FROM projects WHERE id=?", (project_id,)).fetchone()
    project_name = project["name"] if project else "未命名项目"
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
    direction_stats: dict[str, tuple[int, Decimal | None]] = {}
    direction_evidence: dict[str, int | None] = {}
    for direction in _project_directions(conn, project_id):
        aggs = aggregate_project(conn, project_id, direction=direction)
        amounts = [a.cum_amount for a in aggs if a.cum_amount is not None]
        direction_stats[direction] = (
            len(aggs),
            sum(amounts, D(0)) if amounts else None,
        )
        ev = conn.execute(
            """SELECT cr.evidence_id FROM crosscheck_results cr
               JOIN settlement_periods sp ON sp.id=cr.period_id
               WHERE cr.project_id=? AND COALESCE(sp.direction, 'unknown')=?
               ORDER BY cr.checked_at DESC, cr.id DESC LIMIT 1""",
            (project_id, direction),
        ).fetchone()
        direction_evidence[direction] = int(ev["evidence_id"]) if ev and ev["evidence_id"] else None
    high = conn.execute(
        "SELECT COUNT(*) c FROM anomalies WHERE project_id=? AND severity='high'", (project_id,)
    ).fetchone()["c"]
    high_open = conn.execute(
        """SELECT COUNT(*) c FROM anomalies
           WHERE project_id=? AND severity='high' AND status IN ('open', 'deferred')""", (project_id,)
    ).fetchone()["c"]
    supplemented = conn.execute(
        "SELECT COUNT(*) c FROM anomalies WHERE project_id=? AND status='supplemented'", (project_id,)
    ).fetchone()["c"]
    pending_matches = conn.execute(
        "SELECT COUNT(*) c FROM matches WHERE project_id=? AND status='pending'", (project_id,)
    ).fetchone()["c"]
    pending_sheets = pending_sheet_count(conn, project_id)
    incomparable = 0
    for direction in direction_stats:
        incomparable += sum(
            1 for agg in aggregate_project(conn, project_id, direction=direction)
            if agg.status == "incomparable"
        )
    periods = _fetch_periods(conn, project_id)
    n_up = sum(1 for p in periods if p["direction"] == "upward")
    n_down = sum(1 for p in periods if p["direction"] == "downward")
    n_none = len(periods) - n_up - n_down
    n_ev = conn.execute("SELECT COUNT(*) c FROM evidence WHERE project_id=?", (project_id,)).fetchone()["c"]
    level_counts = {r["verification_level"]: int(r["c"]) for r in conn.execute(
        """SELECT verification_level, COUNT(*) c FROM crosscheck_results
           WHERE project_id=? GROUP BY verification_level""", (project_id,)
    )}
    gates = _review_gate_counts(conn, project_id)
    unchecked = max(0, gates["period_count"] - gates["checked_count"])
    gates_open = any((
        not gates["source_files"], not gates["period_count"],
        pending_sheets, high_open, gates["open_findings"], pending_matches,
        level_counts.get("insufficient", 0), level_counts.get("findings", 0),
        gates["range_unproven"], unchecked,
    ))
    if not gates["source_files"]:
        result_status = "尚未导入资料"
    elif not gates["period_count"]:
        result_status = "暂无可审核结算期次"
    else:
        result_status = "审核尚未完成" if gates_open else "当前未发现主要待处理事项"
    doc = docx_lib.Document()
    # python-docx 默认 Title/Normal 使用 Office 主题字体；部分 Mac/WPS 环境会把
    # asciiTheme/majorHAnsi 优先解析为 Calibri，哪怕同时写了东亚字体。显式移除
    # 主题字体引用并写入四套字符字体，避免依赖 Office 专有字体替换。
    theme_attrs = ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme")
    for style_name in ("Normal", "Title", "Default Paragraph Font"):
        style = doc.styles[style_name]
        style.font.name = "Songti SC"
        fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{attr}"), "Songti SC")
        for attr in theme_attrs:
            key = qn(f"w:{attr}")
            if key in fonts.attrib:
                del fonts.attrib[key]

    # python-docx 自带模板的 bestFit 缩放节点可能缺少 w:percent，严格 OOXML
    # 校验器会判为无效。显式补齐，WPS/Word/LibreOffice 均按 100% 处理。
    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None:
        zoom.set(qn("w:percent"), "100")

    # 不使用会重新套用 Office 主题字体的内置 Title 段落；标题运行本身也固定字体。
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(12)
    title_run = title.add_run("CostGuard 管理层摘要")
    title_run.bold = True
    title_run.font.name = "Songti SC"
    title_run.font.size = Pt(26)
    title_fonts = title_run._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        title_fonts.set(qn(f"w:{attr}"), "Songti SC")
    doc.add_paragraph(f"项目名称：{project_name}")
    doc.add_paragraph(f"审核范围：已导入文件 {conn.execute('SELECT COUNT(*) c FROM source_files WHERE project_id=?', (project_id,)).fetchone()['c']} 份，"
                      f"期次 {len(periods)} 期（对上结算 {n_up} 期、对下结算 {n_down} 期、未标记 {n_none} 期）")
    doc.add_paragraph(f"生成时间/版本：{generated_at} / v0.1.7 预览候选")
    status_p = doc.add_paragraph(f"成果状态：{result_status}")
    status_p.runs[0].bold = True

    metrics = doc.add_table(rows=1, cols=2)
    metrics.style = "Table Grid"
    metrics.rows[0].cells[0].text = "关键指标"
    metrics.rows[0].cells[1].text = "当前值"
    metric_rows = [
        ("对上累计", direction_stats.get("upward", (0, None))[1]),
        ("对下累计", direction_stats.get("downward", (0, None))[1]),
        ("期次数", len(periods)),
        ("待确认工作表", pending_sheets),
        ("高风险未处理", high_open),
        ("高风险总数", high),
        ("待补资料异常", supplemented),
        ("不可比清单组", incomparable),
        ("待确认匹配", pending_matches),
        ("校核不充分期数", level_counts.get("insufficient", 0)),
        ("校核有发现期数", level_counts.get("findings", 0)),
        ("尚未校核期数", unchecked),
        ("取数范围未证明工作表", gates["range_unproven"]),
        ("暂不处理异常", gates["deferred"]),
    ]
    for label, value in metric_rows:
        row = metrics.add_row().cells
        row[0].text = label
        if isinstance(value, Decimal):
            row[1].text = f"{round2(value)} 元"
        elif value is None:
            row[1].text = "无法确认"
        else:
            row[1].text = str(value)
    doc.add_paragraph(
        "方向金额仅统计当前可用部分；缺失、待补资料和不可比项目不计入，"
        "不以 0 替代。对应期次校核证据：对上 Evidence ID "
        f"{direction_evidence.get('upward') or '待生成'}；对下 Evidence ID "
        f"{direction_evidence.get('downward') or '待生成'}。"
    )
    doc.add_page_break()

    doc.add_heading("重点结论", level=1)
    conclusion_written = False
    for direction, (group_count, total) in direction_stats.items():
        label = _business_direction(direction)
        total_text = f"{round2(total)} 元" if total is not None else "无法确认"
        ev_text = f"Evidence ID {direction_evidence.get(direction)}" if direction_evidence.get(direction) else "Evidence ID 待生成"
        doc.add_paragraph(
            f"{label}识别清单组 {group_count} 组，累计金额（可用部分）{total_text}；{ev_text}。"
        )
        conclusion_written = True
    if not conclusion_written:
        doc.add_paragraph("当前没有可形成累计金额的清单组，金额无法确认。")
    if level_counts:
        doc.add_paragraph(
            "双向校核分级：" + "；".join(
                f"{'校核充分' if key == 'sufficient' else '校核有发现' if key == 'findings' else '校核不充分'} {value} 期"
                for key, value in sorted(level_counts.items())
            ) + "。校核结果均应按 Evidence ID 回查原始证据。"
        )
    else:
        doc.add_paragraph("尚未生成双向校核结果，暂不能形成校核结论。")

    doc.add_heading("Top 风险事项", level=1)
    risks = conn.execute(
        """SELECT id, rule_id, severity, subject_type, subject_id, evidence_id, message, status
           FROM anomalies WHERE project_id=? AND severity IN ('high', 'medium')
           ORDER BY CASE severity WHEN 'high' THEN 0 WHEN 'medium' THEN 1 ELSE 2 END, id
           LIMIT 10""", (project_id,)
    ).fetchall()
    if risks:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        for cell, text in zip(table.rows[0].cells, ("级别", "风险事项", "处理状态", "金额/结论", "Evidence ID"), strict=False):
            cell.text = text
        for risk in risks:
            cells = table.add_row().cells
            cells[0].text = {"high": "高", "medium": "中", "low": "低"}.get(risk["severity"], "其他")
            cells[1].text = (
                f"{_business_rule(risk['rule_id'])}：{_business_subject(risk['subject_type'])}"
                f"#{risk['subject_id']}；{_normalize_business_text(risk['message'])}"
            )
            cells[2].text = _business_status(risk["status"])
            cells[3].text = "无法确认（需结合原始证据）"
            cells[4].text = f"Evidence ID {risk['evidence_id']}" if risk["evidence_id"] else "待生成"
    else:
        doc.add_paragraph("当前没有可列示的高/中风险事项；这不等同于校核充分。")

    doc.add_heading("待决策/待补资料", level=1)
    pending_lines = []
    if pending_sheets:
        pending_lines.append(f"待确认工作表 {pending_sheets} 张：需人工确认角色、表头、字段和取数范围。")
    if pending_matches:
        pending_lines.append(f"待确认匹配 {pending_matches} 组：非完全匹配需逐项确认并记录原因。")
    if supplemented:
        pending_lines.append(f"已有 {supplemented} 项异常标记为已补资料，请核验补充文件并重新校核。")
    if incomparable:
        pending_lines.append(f"不可比清单组 {incomparable} 组：暂不下金额或责任结论。")
    if gates["range_unproven"]:
        pending_lines.append(
            f"有 {gates['range_unproven']} 张工作表取数范围完整性无法证明：需人工核对原表边界。"
        )
    if unchecked:
        pending_lines.append(f"尚未执行双向校核 {unchecked} 期：暂不能形成充分校核结论。")
    if gates["deferred"]:
        pending_lines.append(f"暂不处理异常 {gates['deferred']} 项：保留原状态，不视为已闭合。")
    if not pending_lines:
        pending_lines.append("当前没有登记的待决策事项；仍需按 Evidence ID 完成人工复核和业务审批。")
    for line in pending_lines:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("审核范围与限制", level=1)
    doc.add_paragraph(
        f"统计范围：期次 {len(periods)} 期（对上结算 {n_up} 期、对下结算 {n_down} 期、未标记方向 {n_none} 期）；"
        "对上、对下和未标记方向未作跨方向净额或合计。"
    )
    doc.add_paragraph("数据纪律：缺失数据未自动补 0；不可比数据未强行比较；金额无法确认时明确写“无法确认”。")
    doc.add_paragraph(
        f"校核门控：已执行 {gates['checked_count']} / {gates['period_count']} 期；"
        f"校核不充分 {gates['insufficient']} 期；校核有发现 {gates['findings']} 期；"
        f"取数范围未证明 {gates['range_unproven']} 张工作表。"
    )
    doc.add_paragraph(
        "追溯说明：本 Word 摘要不包含证据入口按钮；正文管理结论均标注 Evidence ID，"
        "待生成或无法确认处不会伪造证据。逐单元格证据链与证据索引"
        f"（当前共 {n_ev} 条证据记录）请查阅 CostGuard 导出的 Excel 审核底稿工作簿"
        "《证据索引》工作表，按证据 ID 查询。"
    )
    doc.add_paragraph(
        "本摘要由软件自动生成，仅反映已导入数据；自动计算结果必须经过人工复核和业务审批。"
        "未经批准，不构成任何真实业务结论，包括最终结算、责任认定或正式管理结论。"
    )
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"CostGuard管理层摘要_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(str(path))
    return path
