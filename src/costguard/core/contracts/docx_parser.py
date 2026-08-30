"""合同文件文本解析：docx / pdf / txt。

产出统一结构：段落列表 [{index, text}]（index 为 1-based 段落号/页码标记）。
纪律：只读；解析失败显式报告。
"""
from __future__ import annotations

from pathlib import Path


def parse_docx(path: Path) -> list[dict]:
    import docx  # python-docx

    doc = docx.Document(str(path))
    paras = []
    for i, p in enumerate(doc.paragraphs, start=1):
        text = (p.text or "").strip()
        if text:
            paras.append({"index": i, "text": text})
    # 表格文本也纳入（合同常用表格约定工期/金额）
    for ti, table in enumerate(doc.tables, start=1):
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paras.append({"index": f"t{ti}", "text": " | ".join(cells)})
    return paras


def parse_pdf(path: Path) -> list[dict]:
    import pdfplumber

    paras = []
    with pdfplumber.open(str(path)) as pdf:
        for pi, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            for j, line in enumerate(text.splitlines(), start=1):
                line = line.strip()
                if line:
                    paras.append({"index": f"p{pi}:{j}", "text": line})
    if not paras:
        # 扫描件无文本层时必须显式失败（验收执行器/Gui 均有 NotImplementedError
        # 通道），绝不静默返回空结果让用户误以为"合同里没有该条款"。
        raise NotImplementedError(
            f"PDF 未提取到任何文本层（疑似扫描件，OCR 尚未支持）：{path.name}；"
            "扫描资料不得直接形成业务结论，需人工复核原件"
        )
    return paras


def parse_txt(path: Path) -> list[dict]:
    text = Path(path).read_text(encoding="utf-8", errors="replace")
    return [{"index": i, "text": t.strip()} for i, t in enumerate(text.splitlines(), start=1) if t.strip()]


def parse_contract(path: Path, file_type: str) -> list[dict]:
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(str(path))
    if file_type == "docx":
        return parse_docx(path)
    if file_type == "pdf":
        return parse_pdf(path)
    if file_type == "txt":
        return parse_txt(path)
    if file_type in ("doc", "image"):
        raise NotImplementedError(f"parser for '{file_type}' not implemented yet")
    raise ValueError(f"unsupported contract file type: {file_type}")
