"""合同文件文本解析：docx / pdf / txt。

产出统一结构：段落列表 [{index, text}]（index 为 1-based 段落号/页码标记）。
纪律：只读；解析失败显式报告。
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any

from jiadun.core.parsing.pdf_pipeline import (
    OcrProvider,
    PdfExtractionReport,
    PdfRenderer,
    extract_pdf_document,
    paragraphs_from_report,
)


def _ensure_magic(path: Path, magics: tuple[bytes, ...], label: str) -> None:
    """魔数前置校验：同步残留/损坏/误命名文件显式拒绝，不给解析器抛裸异常。"""
    with open(path, "rb") as fh:
        head = fh.read(8)
    if not any(head.startswith(m) for m in magics):
        raise ValueError(
            f"不是有效的 {label} 文件（文件头不符，可能是同步残留、损坏或误命名文件）：{path.name}"
        )


def parse_docx(path: Path) -> list[dict]:
    import docx  # python-docx

    doc = docx.Document(str(path))
    paras = []
    for i, p in enumerate(doc.paragraphs, start=1):
        text = (p.text or "").strip()
        if text:
            paras.append({"index": i, "text": text})
    # 表格文本也纳入（合同常用表格约定工期/金额）
    for ti, table in enumerate(doc.tables, start=1):
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paras.append({"index": f"t{ti}", "text": " | ".join(cells)})
    return paras


@dataclass(frozen=True, slots=True)
class ContractParseResult:
    """兼容段落结果之外，向导入层提供 PDF 页级提取快照。"""

    paragraphs: list[dict[str, Any]]
    pdf_report: PdfExtractionReport | None = None


def parse_pdf(
    path: Path,
    *,
    renderer: PdfRenderer | None = None,
    ocr_provider: OcrProvider | None = None,
) -> list[dict[str, Any]]:
    """逐页解析 PDF；页面未完整解释时抛出兼容的专用待处理异常。"""
    report = extract_pdf_document(
        Path(path), renderer=renderer, ocr_provider=ocr_provider
    )
    return paragraphs_from_report(report)


def parse_txt(path: Path) -> list[dict]:
    text = Path(path).read_text(encoding="utf-8", errors="replace")
    return [{"index": i, "text": t.strip()} for i, t in enumerate(text.splitlines(), start=1) if t.strip()]


def parse_contract_result(
    path: Path,
    file_type: str,
    *,
    renderer: PdfRenderer | None = None,
    ocr_provider: OcrProvider | None = None,
) -> ContractParseResult:
    """解析合同文本，并在 PDF 时返回页级提取报告。"""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(str(path))
    if file_type == "docx":
        _ensure_magic(path, (b"PK",), "DOCX")
        return ContractParseResult(parse_docx(path))
    if file_type == "pdf":
        _ensure_magic(path, (b"%PDF-",), "PDF")
        report = extract_pdf_document(
            path, renderer=renderer, ocr_provider=ocr_provider
        )
        return ContractParseResult(paragraphs_from_report(report), report)
    if file_type == "txt":
        return ContractParseResult(parse_txt(path))
    if file_type in ("doc", "image"):
        raise NotImplementedError(f"parser for '{file_type}' not implemented yet")
    raise ValueError(f"unsupported contract file type: {file_type}")


def parse_contract(
    path: Path,
    file_type: str,
    *,
    renderer: PdfRenderer | None = None,
    ocr_provider: OcrProvider | None = None,
) -> list[dict[str, Any]]:
    """保留旧返回类型的合同解析兼容入口。"""
    return parse_contract_result(
        path, file_type, renderer=renderer, ocr_provider=ocr_provider
    ).paragraphs
