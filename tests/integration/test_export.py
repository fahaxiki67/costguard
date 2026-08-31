"""导出集成测试：从合成数据到完整 Excel 底稿（含公式回读验证）。"""
import sys
from decimal import Decimal
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).parents[2] / "synthetic_test_data"))

from generator import make_messy, make_multi_period  # noqa: E402

from costguard.core.anomalies import engine  # noqa: E402
from costguard.core.contracts import run_contract  # noqa: E402
from costguard.core.engine import settlement_io  # noqa: E402
from costguard.core.engine.money import round2  # noqa: E402
from costguard.core.export import excel_export  # noqa: E402


@pytest.fixture()
def full_project(tmp_path):
    from costguard.core.models import project as pm

    info = pm.create_project("导出-全流程", tmp_path / "ws")
    info, conn = pm.open_project(Path(info.workspace_path))
    pdir = Path(info.workspace_path)
    src1 = pdir.parent / "multi.xlsx"
    make_multi_period(src1, periods=3)
    settlement_io.import_settlement_file(conn, info.project_id, pdir, src1)
    src2 = pdir.parent / "messy.xlsx"
    make_messy(src2, seed=7)
    settlement_io.import_settlement_file(conn, info.project_id, pdir, src2)
    engine.run_anomalies(conn, info.project_id)
    yield info, conn, pdir / "exports"
    conn.close()



class TestExcelExport:
    @pytest.mark.parametrize("export_kind", ["excel_workbook", "management_summary_docx"])
    def test_fail_closed_state_blocks_current_export_registration(
        self, full_project, export_kind
    ):
        """运行级不可用时，Excel/Word 均不得生成或登记为当前成果。"""
        info, conn, exports = full_project
        run_contract.set_fail_closed_state(
            conn, info.project_id, reason="synthetic database is not writable"
        )
        exporter = (
            excel_export.export_workbook
            if export_kind == "excel_workbook"
            else excel_export.export_management_summary_docx
        )
        with pytest.raises(run_contract.CurrentResultsUnavailableError, match="当前结果不可用"):
            exporter(conn, info.project_id, exports)

        assert not list(exports.glob("CostGuard审核底稿_*.xlsx"))
        assert not list(exports.glob("CostGuard管理层摘要_*.docx"))
        assert all(
            item["status"] != "current"
            for item in run_contract.export_status(conn, info.project_id, export_kind)
        )

    def test_amount_only_recompute_is_visible_in_period_and_comparison_outputs(
        self, tmp_path
    ):
        from tests.integration.test_aggregate_crosscheck import _make_amount_case

        info, conn, period_id = _make_amount_case(tmp_path)
        try:
            path = excel_export.export_workbook(conn, info.project_id, tmp_path / "exports")
            import openpyxl

            wb = openpyxl.load_workbook(path, data_only=False)
            summary = wb["对下结算累计表"]
            assert summary.cell(row=2, column=4).value == Decimal("200")
            amount_diff = wb["金额差异表"]
            assert any(
                amount_diff.cell(row=r, column=5).value == Decimal("200")
                for r in range(2, amount_diff.max_row + 1)
            )
            comparison = wb["对上对下对比表"]
            assert comparison.max_row == 2
            assert comparison.cell(row=2, column=6).value == Decimal("200")
            assert "待补资料" in str(comparison.cell(row=2, column=8).value)
            assert conn.execute(
                "SELECT amount_sum FROM period_totals WHERE period_id=?",
                (period_id,),
            ).fetchone() is None
        finally:
            conn.close()
    def test_exported_tables_have_borders_and_centered_cells(self, full_project):
        """正式导出的各张表应具备可直接验收的基础表格格式。"""
        info, conn, exports = full_project
        path = excel_export.export_workbook(conn, info.project_id, exports)
        import openpyxl

        wb = openpyxl.load_workbook(path, data_only=False)
        for ws in wb.worksheets:
            assert ws.max_row >= 1 and ws.max_column >= 1
            for row in ws.iter_rows(
                min_row=1, max_row=ws.max_row, min_col=1, max_col=ws.max_column
            ):
                for cell in row:
                    assert cell.border.left.style == "thin", (ws.title, cell.coordinate)
                    assert cell.border.right.style == "thin", (ws.title, cell.coordinate)
                    assert cell.border.top.style == "thin", (ws.title, cell.coordinate)
                    assert cell.border.bottom.style == "thin", (ws.title, cell.coordinate)
                    assert cell.alignment.horizontal == "center", (ws.title, cell.coordinate)
                    assert cell.alignment.vertical == "center", (ws.title, cell.coordinate)

    def test_workbook_created_all_sheets(self, full_project):
        info, conn, exports = full_project
        path = excel_export.export_workbook(conn, info.project_id, exports)
        assert path.exists()
        import openpyxl

        wb = openpyxl.load_workbook(path)
        expected = {"管理层摘要", "未标记累计表", "对上对下对比表", "单价差异表", "工程量差异表",
                    "金额差异表", "异常清单", "待核实事项清单", "证据索引", "审核底稿"}
        assert expected <= set(wb.sheetnames), wb.sheetnames

    def test_audit_worksheet_has_formulas(self, full_project):
        info, conn, exports = full_project
        path = excel_export.export_workbook(conn, info.project_id, exports)
        import openpyxl

        wb = openpyxl.load_workbook(path)
        ws = wb["审核底稿"]
        formulas = [ws.cell(row=r, column=7).value for r in range(2, ws.max_row + 1)]
        assert any(isinstance(v, str) and v.startswith("=E") for v in formulas), formulas[:5]
        program_vals = [ws.cell(row=r, column=8).value for r in range(2, ws.max_row + 1)]
        assert any(isinstance(v, (int, float)) for v in program_vals), (
            "程序计算合价列必须有 Decimal 数值（不依赖 Office 重算）")
        diffs = [ws.cell(row=r, column=10).value for r in range(2, ws.max_row + 1)]
        assert any(isinstance(v, str) and v.startswith("=ROUND(") for v in diffs)

    def test_formula_references_match_db(self, full_project):
        """公式引用正确性：每行 =E{r}*F{r} 引用的单元格值必须与 DB 该行数量/单价一致。

        这是公式可复算的静态保障——行号错位会让公式"能算"但算错。
        """
        info, conn, exports = full_project
        path = excel_export.export_workbook(conn, info.project_id, exports)
        import openpyxl

        wb = openpyxl.load_workbook(path, data_only=False)
        ws = wb["审核底稿"]
        db_rows = {
            r["id"]: (r["code"], r["name"], r["quantity"], r["unit_price"])
            for r in conn.execute(
                """SELECT li.id, li.code, li.name, li.quantity, li.unit_price FROM line_items li
                   JOIN settlement_periods sp ON sp.id = li.period_id
                   WHERE sp.project_id=? AND li.flags_json NOT LIKE '%"subtotal": true%'""",
                (info.project_id,),
            )
        }
        checked = 0
        for r in range(2, ws.max_row + 1):
            formula = ws.cell(row=r, column=7).value
            if not (isinstance(formula, str) and formula.startswith("=E")):
                continue
            assert formula == f"=E{r}*F{r}", f"formula row mismatch at {r}: {formula}"
            row_id = ws.cell(row=r, column=14).value  # 行ID
            assert row_id in db_rows, f"worksheet row {r} references unknown line_item"
            d_code, d_name, d_qty, d_price = db_rows[row_id]
            assert ws.cell(row=r, column=2).value == d_code
            assert ws.cell(row=r, column=3).value == d_name
            qty = ws.cell(row=r, column=5).value
            price = ws.cell(row=r, column=6).value
            assert qty is not None and float(qty) == float(d_qty), f"qty mismatch at row {r}"
            assert price is not None and float(price) == float(d_price), f"price mismatch at row {r}"
            checked += 1
        assert checked > 0

    def test_formulas_recalculate_in_real_engine(self, full_project, tmp_path):
        """真实计算引擎重算（LibreOffice，与 WPS 同遵 OOXML 公式规范）：

        重算后 G列=数量×单价、I列=ROUND(底稿合价-原表合价,2) 必须与 DB 值一致。
        """
        import os
        import shutil
        import subprocess

        # 优先真二进制（Homebrew /opt/homebrew/bin/soffice 是 2 行 sh 包装脚本，
        # 偶发 SIGABRT 发生在其 exec 瞬间——命令矩阵证明两种入口转换结果等效）
        real_bin = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
        soffice = str(real_bin) if real_bin.exists() else shutil.which("soffice")
        if soffice is None:
            pytest.skip("LibreOffice not installed")

        info, conn, exports = full_project
        path = excel_export.export_workbook(conn, info.project_id, exports)
        out = tmp_path / "recalc"
        out.mkdir(parents=True, exist_ok=True)  # 显式创建输出目录（合法前置条件）
        # 隔离 UserInstallation：独立 profile，不影响任何其他应用
        profile = tmp_path / "lo_profile_isolated"
        profile.mkdir(parents=True, exist_ok=True)
        cmd = [soffice,
               "-env:UserInstallation=" + Path(profile).as_uri(),
               "--headless", "--convert-to", "xlsx", "--outdir", str(out), str(path)]

        # 根因（崩溃报告 soffice-*.ips，4 份同栈）：headless 启动器早期初始化 AppKit，
        # RegisterApplication（HIServices）与驻留 LibreOffice GUI 实例的菜单栏状态
        # 竞态 → SIGABRT。防御组合：
        # 1) 每次尝试以唯一 __CFBundleIdentifier 注册，避免与 GUI 实例身份冲突（针对根因）；
        # 2) 指数退避 0/1/2s（崩溃报告显示零退避的连续重试全撞同一竞态窗口）；
        # 3) 仅 SIGABRT（rc 134/-6）重试，其余 rc 立即失败；逐次留 rc/stdout/stderr 证据；
        # 4) 转换成功后的全部重算断言不重试、不放宽。
        max_attempts = 3
        attempt_log: list[str] = []
        proc = None
        for attempt in range(1, max_attempts + 1):
            if attempt > 1:
                import time

                time.sleep(attempt - 1)  # 退避 1s / 2s，穿过竞态窗口
            env = dict(os.environ)
            env["__CFBundleIdentifier"] = f"org.costguard.headless.{os.getpid()}.{attempt}"
            proc = subprocess.run(cmd, capture_output=True, text=True, timeout=240,
                                  env=env)
            if proc.returncode == 0:
                break
            attempt_log.append(
                f"attempt{attempt} rc={proc.returncode} stderr={proc.stderr[-300:]}")
            if proc.returncode == 1 and not proc.stdout and not proc.stderr:
                pytest.skip(
                    "LibreOffice headless conversion is unavailable in this macOS "
                    "test environment (rc=1 with no output); keep WPS/manual gate "
                    "separate from the application test suite."
                )
            if proc.returncode not in (134, -6):
                pytest.fail(
                    f"soffice 非启动器崩溃错误 rc={proc.returncode}\n"
                    f"cmd={cmd}\nstdout(尾2000)={proc.stdout[-2000:]}\n"
                    f"stderr(尾2000)={proc.stderr[-2000:]}"
                )
        else:
            # 环境自检：SIGABRT 耗尽时附带安装完整性证据（LibreOffice 安装损坏时
            # 密封校验会失败，这是启动器偶发 SIGABRT 的常见根源）
            try:
                cs = subprocess.run(
                    ["codesign", "--verify", "--deep", "--strict",
                     "/Applications/LibreOffice.app"],
                    capture_output=True, text=True, timeout=120)
                cs_status = "安装完整性 OK" if cs.returncode == 0 else (
                    f"安装损坏（codesign rc={cs.returncode}: {cs.stdout.strip()[:200]}"
                    f"{cs.stderr.strip()[:200]}）—— 请重装 LibreOffice 或改用专用验证机")
            except Exception as exc:  # noqa: BLE001
                cs_status = f"codesign 自检失败: {exc}"
            pytest.fail(
                f"soffice 启动器连续 {max_attempts} 次 SIGABRT（环境阻断证据）\n"
                + "\n".join(attempt_log)
                + f"\n环境自检: {cs_status}"
            )
        recalc_path = out / path.name
        assert recalc_path.exists(), (
            f"转换未产出文件: stdout={proc.stdout[-500:] if proc else ''} log={attempt_log}")

        from decimal import Decimal

        import openpyxl

        wb = openpyxl.load_workbook(recalc_path, data_only=True)
        ws = wb["审核底稿"]
        db_amounts = {
            r["id"]: r["amount"]
            for r in conn.execute(
                """SELECT li.id, li.amount FROM line_items li
                   JOIN settlement_periods sp ON sp.id = li.period_id
                   WHERE sp.project_id=? AND li.flags_json NOT LIKE '%"subtotal": true%'""",
                (info.project_id,),
            )
        }
        checked = 0
        for r in range(2, ws.max_row + 1):
            qty = ws.cell(row=r, column=5).value
            price = ws.cell(row=r, column=6).value
            g = ws.cell(row=r, column=7).value  # 底稿合价（公式重算值，未舍入的精确乘积）
            h_prog = ws.cell(row=r, column=8).value  # 程序计算合价（写入的 Decimal 值）
            h_orig = ws.cell(row=r, column=9).value  # 原表合价
            if qty is None or price is None or g is None:
                continue
            # 勾稽1：G 必须等于 数量×单价 的精确乘积（数量可能除不尽，不预设舍入口径）
            expect_exact = Decimal(repr(qty)) * Decimal(repr(price))
            assert abs(Decimal(repr(g)) - expect_exact) <= Decimal("0.001"), \
                f"recalc mismatch row {r}: G={g} expect={expect_exact}"
            # 勾稽1b（双值）：G 重算值必须等于程序写入的 Decimal 值
            # ——主结论不依赖 Office 重算的直接证据
            if h_prog is not None and not isinstance(h_prog, str):
                assert abs(Decimal(repr(g)) - Decimal(repr(h_prog))) <= Decimal("0.001"), \
                    f"program value mismatch row {r}: G={g} program={h_prog}"
            # 勾稽2：差异列 J == ROUND(G-I, 2)（同源数据时必为 0）
            j_val = ws.cell(row=r, column=10).value
            if j_val is not None and h_orig is not None:
                expect_diff = round2(Decimal(repr(g)) - Decimal(repr(h_orig)))
                assert abs(Decimal(repr(j_val)) - expect_diff) <= Decimal("0.001"), \
                    f"diff col mismatch row {r}: J={j_val} expect={expect_diff}"
            # 勾稽3：原表合价与 DB 一致（LibreOffice 缓存值为 15 位有效数字，
            # 允许 1e-6 的表示误差——远小于业务容差 0.02）
            row_id = ws.cell(row=r, column=14).value
            if row_id in db_amounts and db_amounts[row_id] is not None:
                assert abs(Decimal(repr(h_orig)) - Decimal(db_amounts[row_id])) <= Decimal("1e-6"), \
                    f"original amount mismatch row {r}: {h_orig} vs {db_amounts[row_id]}"
            checked += 1
        assert checked > 0

    def test_ooxml_structure_wps_compatible(self, full_project):
        """WPS/OOXML 兼容性静态检查：
        - 有效 zip 容器且含 [Content_Types].xml；
        - sheet 名 ≤31 字符且不含非法字符；
        - 金额列为标准数字格式（#,##0.00）。
        """
        import zipfile

        info, conn, exports = full_project
        path = excel_export.export_workbook(conn, info.project_id, exports)
        with zipfile.ZipFile(path) as z:
            names = z.namelist()
            assert "[Content_Types].xml" in names
            assert "xl/workbook.xml" in names
            assert z.testzip() is None

        import openpyxl

        wb = openpyxl.load_workbook(path)
        for name in wb.sheetnames:
            assert len(name) <= 31, f"sheet name too long: {name}"
            assert not any(ch in name for ch in r"[]:*?/\\"), f"illegal chars in sheet name: {name}"
        ws = wb["审核底稿"]
        fmts = {ws.cell(row=r, column=8).number_format for r in range(2, min(ws.max_row + 1, 6))}
        assert all(f == excel_export.MONEY_FMT for f in fmts), fmts

    def test_summary_and_anomaly_lists_consistent(self, full_project):
        info, conn, exports = full_project
        path = excel_export.export_workbook(conn, info.project_id, exports)
        import openpyxl

        wb = openpyxl.load_workbook(path)
        n_high = conn.execute(
            "SELECT COUNT(*) c FROM anomalies WHERE project_id=? AND severity='high'", (info.project_id,)
        ).fetchone()["c"]
        rows = list(wb["异常清单"].iter_rows(min_row=2, values_only=True))
        assert sum(1 for r in rows if r[3] == "高") == n_high
        ev_rows = list(wb["证据索引"].iter_rows(min_row=2, values_only=True))
        n_ev = conn.execute("SELECT COUNT(*) c FROM evidence WHERE project_id=?", (info.project_id,)).fetchone()["c"]
        assert len(ev_rows) == n_ev

    def test_updown_comparison_marks_missing_side(self, full_project):
        """方向未标记 → 明确写"不可比"，不得强行比较；一侧缺失 → 待补资料。"""
        info, conn, exports = full_project
        path = excel_export.export_workbook(conn, info.project_id, exports)
        import openpyxl

        wb = openpyxl.load_workbook(path)
        ws = wb["对上对下对比表"]
        rows = list(ws.iter_rows(min_row=2, values_only=True))
        assert rows and "无可比数据" in str(rows[0][1])

        # 标记方向后：对上/对下各自累计，同清单差异公式生成
        with conn:
            conn.execute("UPDATE settlement_periods SET direction='upward' WHERE period_no<=2")
            conn.execute("UPDATE settlement_periods SET direction='downward' WHERE period_no>2")
        path2 = excel_export.export_workbook(conn, info.project_id, exports)
        wb2 = openpyxl.load_workbook(path2)
        ws2 = wb2["对上对下对比表"]
        formulas = [ws2.cell(row=r, column=7).value for r in range(2, ws2.max_row + 1)]
        assert any(isinstance(v, str) and v.startswith("=E") for v in formulas)
        notes = [str(ws2.cell(row=r, column=8).value) for r in range(2, ws2.max_row + 1)]
        # C25 清单对上/对下名称不同（生成器改名）→ 必须提示核实，不得静默合并
        assert any("名称不一致" in v for v in notes), notes

        # 单侧缺失场景：插入一个仅存在于对下方向的清单 → 必须标"待补资料"，不得按 0 比较
        with conn:
            p4 = conn.execute(
                "SELECT id FROM settlement_periods WHERE project_id=? AND period_no=4",
                (info.project_id,),
            ).fetchone()
            assert p4
            conn.execute(
                "INSERT INTO line_items(period_id, code, name, unit, quantity, amount)"
                " VALUES (?, 'XTEST', '仅对下清单', 'm3', '10', '1000')",
                (p4["id"],),
            )
        path3 = excel_export.export_workbook(conn, info.project_id, exports)
        wb3 = openpyxl.load_workbook(path3)
        ws3 = wb3["对上对下对比表"]
        for r in range(2, ws3.max_row + 1):
            if ws3.cell(row=r, column=2).value == "仅对下清单":
                assert ws3.cell(row=r, column=7).value is None, "missing side must not be zero-filled"
                assert "待补资料" in str(ws3.cell(row=r, column=8).value)
                break
        else:
            pytest.fail("down-only item not found in comparison sheet")

    def test_docx_summary(self, full_project):
        info, conn, exports = full_project
        path = excel_export.export_management_summary_docx(conn, info.project_id, exports)
        assert path.exists()
        import docx as docx_lib

        doc = docx_lib.Document(str(path))
        assert "CostGuard 管理层摘要" in doc.paragraphs[0].text
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "合成测试数据" not in text, "通用成果不得把真实资料误写成合成数据"
        assert "人工复核" in text and "业务审批" in text
        from docx.oxml.ns import qn

        for style_name in ("Normal", "Title"):
            style = doc.styles[style_name]
            assert style.font.name == "Songti SC"
            assert style._element.rPr.rFonts.get(qn("w:eastAsia")) == "Songti SC"
            assert style._element.rPr.rFonts.get(qn("w:asciiTheme")) is None
            assert style._element.rPr.rFonts.get(qn("w:hAnsiTheme")) is None
        assert doc.paragraphs[0].style.name == "Normal"
        title_run = doc.paragraphs[0].runs[0]
        assert title_run.font.name == "Songti SC"
        assert title_run._element.rPr.rFonts.get(qn("w:eastAsia")) == "Songti SC"
        zoom = doc.settings.element.find(qn("w:zoom"))
        assert zoom is not None and zoom.get(qn("w:percent")) == "100"

    def test_workbook_disclaimer_is_valid_for_real_or_synthetic_data(self, full_project):
        info, conn, exports = full_project
        path = excel_export.export_workbook(conn, info.project_id, exports)
        import openpyxl

        wb = openpyxl.load_workbook(path, data_only=False)
        text = "\n".join(
            str(cell.value) for row in wb["管理层摘要"].iter_rows() for cell in row
            if cell.value is not None
        )
        assert "合成测试数据" not in text
        assert "人工复核" in text and "业务审批" in text
