"""对上控制基准候选与上限比较测试（宪章 §六）。

候选只能来自已确认事实或人工显式登记；确认必须给依据；比较输出固定五态；
两个并存有效基准必须 CONTROL_CONFLICT，不允许自动挑选。
"""

import docx as docx_lib
import pytest
from tests.unit.test_contract_extract import SAMPLE_CONTRACT

from jiadun.core.contracts import extract
from jiadun.core.db import migrations
from jiadun.core.engine import control_baseline as cb


@pytest.fixture()
def project_db(tmp_path):
    db_path = tmp_path / "project.db"
    migrations.migrate(db_path, tmp_path / "backups")
    conn = migrations.connect(db_path)
    with conn:
        project_id = conn.execute(
            """INSERT INTO projects(name, schema_version, workspace_path, created_at)
               VALUES (?,?,?,?)""",
            ("控制基准测试", migrations.LATEST_SCHEMA_VERSION, str(tmp_path), "2026"),
        ).lastrowid
    yield conn, int(project_id), tmp_path
    conn.close()


@pytest.fixture()
def confirmed_amount_fact(project_db):
    """导入合成合同并把合同价款事实确认为 confirmed。"""
    conn, pid, pdir = project_db
    src = pdir / "分包合同.docx"
    d = docx_lib.Document()
    for line in SAMPLE_CONTRACT.splitlines():
        d.add_paragraph(line)
    d.save(str(src))
    extract.import_contract(conn, pid, pdir, src, doc_type="subcontract")
    fact_ids = [
        int(r["id"]) for r in conn.execute(
            """SELECT cf.id FROM contract_facts cf
               JOIN contract_docs cd ON cd.id=cf.doc_id
               WHERE cd.project_id=? AND cf.fact_key='contract_amount'
                 AND cf.fact_value LIKE '%1286500%'""",
            (pid,),
        ).fetchall()
    ]
    assert fact_ids
    extract.set_fact_review(conn, pid, fact_ids[0], "confirmed", reason="与合同原文核对一致")
    return conn, pid, fact_ids[0]


class TestCandidateLifecycle:
    def test_unconfirmed_fact_cannot_become_candidate(self, project_db):
        conn, pid, pdir = project_db
        src = pdir / "分包合同.docx"
        d = docx_lib.Document()
        for line in SAMPLE_CONTRACT.splitlines():
            d.add_paragraph(line)
        d.save(str(src))
        extract.import_contract(conn, pid, pdir, src, doc_type="subcontract")
        fact_id = int(conn.execute(
            """SELECT cf.id FROM contract_facts cf
               JOIN contract_docs cd ON cd.id=cf.doc_id
               WHERE cd.project_id=? AND cf.fact_key='contract_amount' LIMIT 1""",
            (pid,),
        ).fetchone()["id"])
        with pytest.raises(ValueError, match="已确认"):
            cb.create_candidate_from_fact(conn, pid, fact_id)

    def test_candidate_from_confirmed_fact(self, confirmed_amount_fact):
        conn, pid, fact_id = confirmed_amount_fact
        cb.create_candidate_from_fact(conn, pid, fact_id)
        baselines = cb.list_baselines(conn, pid)
        assert len(baselines) == 1
        b = baselines[0]
        assert b["status"] == "candidate"
        assert b["amount"].startswith("1286500")
        assert b["doc_title"] == "分包合同"

    def test_manual_candidate_requires_source(self, project_db):
        conn, pid, _ = project_db
        with pytest.raises(ValueError, match="出处"):
            cb.create_candidate_manual(conn, pid, "1286500", source_note="  ")
        with pytest.raises(ValueError, match="金额"):
            cb.create_candidate_manual(conn, pid, "abc", source_note="审计报告审定金额")

    def test_confirm_requires_reason(self, confirmed_amount_fact):
        conn, pid, fact_id = confirmed_amount_fact
        baseline_id = cb.create_candidate_from_fact(conn, pid, fact_id)
        with pytest.raises(ValueError, match="核对依据"):
            cb.set_baseline_review(conn, pid, baseline_id, "confirmed")
        cb.set_baseline_review(
            conn, pid, baseline_id, "confirmed",
            reason="终审报告审定表已核对，含税口径一致",
        )
        b = cb.list_baselines(conn, pid)[0]
        assert b["status"] == "confirmed"
        assert b["confirmed_by"] == "user"


class TestCompare:
    def test_unconfirmed_baseline_is_pending(self, confirmed_amount_fact):
        conn, pid, fact_id = confirmed_amount_fact
        baseline_id = cb.create_candidate_from_fact(conn, pid, fact_id)
        result = cb.compare_upward_result(conn, pid, baseline_id, "1200000")
        assert result["status"] == "PENDING"

    def test_unknown_tax_basis_is_pending(self, confirmed_amount_fact):
        conn, pid, fact_id = confirmed_amount_fact
        baseline_id = cb.create_candidate_from_fact(conn, pid, fact_id)
        cb.set_baseline_review(conn, pid, baseline_id, "confirmed", reason="核对一致")
        result = cb.compare_upward_result(conn, pid, baseline_id, "1200000")
        assert result["status"] == "PENDING"
        assert "税口径" in result["reason"]

    def test_tax_mismatch_is_incomparable(self, confirmed_amount_fact):
        conn, pid, fact_id = confirmed_amount_fact
        baseline_id = cb.create_candidate_from_fact(conn, pid, fact_id, tax_basis="included")
        cb.set_baseline_review(conn, pid, baseline_id, "confirmed", reason="核对一致")
        result = cb.compare_upward_result(
            conn, pid, baseline_id, "1200000", settlement_tax_basis="excluded"
        )
        assert result["status"] == "INCOMPARABLE"

    def test_pass_and_fail_with_delta(self, confirmed_amount_fact):
        conn, pid, fact_id = confirmed_amount_fact
        baseline_id = cb.create_candidate_from_fact(conn, pid, fact_id, tax_basis="included")
        cb.set_baseline_review(conn, pid, baseline_id, "confirmed", reason="核对一致")
        ok = cb.compare_upward_result(
            conn, pid, baseline_id, "1200000.00", settlement_tax_basis="included"
        )
        assert ok["status"] == "PASS"
        assert ok["delta"] == "-86500.00"
        over = cb.compare_upward_result(
            conn, pid, baseline_id, "1300000.00", settlement_tax_basis="included"
        )
        assert over["status"] == "FAIL"
        assert over["delta"] == "13500.00"

    def test_conflict_without_supersedes(self, confirmed_amount_fact):
        conn, pid, fact_id = confirmed_amount_fact
        first = cb.create_candidate_from_fact(conn, pid, fact_id, tax_basis="included")
        cb.set_baseline_review(conn, pid, first, "confirmed", reason="初审报告核对一致")
        second = cb.create_candidate_from_fact(
            conn, pid, fact_id, tax_basis="included",
        )
        cb.set_baseline_review(conn, pid, second, "confirmed", reason="终审报告核对一致")
        result = cb.compare_upward_result(
            conn, pid, second, "1200000", settlement_tax_basis="included"
        )
        assert result["status"] == "CONTROL_CONFLICT"

        # 冲突解除：终审版显式声明取代初审版 + 人工拒绝初审候选
        cb.set_baseline_review(conn, pid, first, "rejected", reason="被终审报告取代")
        third = cb.create_candidate_from_fact(
            conn, pid, fact_id, tax_basis="included", supersedes_id=second
        )
        cb.set_baseline_review(conn, pid, third, "confirmed", reason="终审版取代初审版")
        active = cb.compare_upward_result(
            conn, pid, third, "1200000", settlement_tax_basis="included"
        )
        assert active["status"] == "PASS"


class TestUpwardPeriods:
    def test_upward_periods_with_decimal_totals(self, project_db):
        conn, pid, _ = project_db
        with conn:
            cur = conn.execute(
                """INSERT INTO settlement_periods(project_id, period_no, title, direction, tax_mode)
                   VALUES (?,?,?,?,?)""",
                (pid, 1, "第一期对上结算", "upward", "included"),
            )
            period_id = cur.lastrowid
            for amount in ("1000.50", "2000.25", None):
                conn.execute(
                    """INSERT INTO line_items(period_id, name, amount) VALUES (?,?,?)""",
                    (period_id, "明细项", amount),
                )
            conn.execute(
                """INSERT INTO settlement_periods(project_id, period_no, title, direction)
                   VALUES (?,?,?,?)""",
                (pid, 2, "对下分包结算", "downward"),
            )
        periods = cb.list_upward_periods(conn, pid)
        assert len(periods) == 1
        assert periods[0]["amount_total"] == "3000.75"
        assert periods[0]["detail_rows"] == 3


class TestEvidence:
    def test_all_three_kinds_written(self, confirmed_amount_fact):
        conn, pid, fact_id = confirmed_amount_fact
        baseline_id = cb.create_candidate_from_fact(conn, pid, fact_id)
        cb.set_baseline_review(conn, pid, baseline_id, "confirmed", reason="核对一致")
        cb.compare_upward_result(conn, pid, baseline_id, "100", settlement_tax_basis="included")
        kinds = {
            r["kind"] for r in conn.execute(
                "SELECT DISTINCT kind FROM evidence WHERE project_id=?", (pid,)
            ).fetchall()
        }
        assert {"control_baseline", "control_baseline_review", "control_baseline_compare"} <= kinds
