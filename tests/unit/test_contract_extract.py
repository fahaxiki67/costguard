"""合同条款提取测试：原文引用必须随身携带。"""

import json

import pytest

from jiadun.core.contracts import docx_parser, extract, run_contract
from jiadun.core.db import migrations

SAMPLE_CONTRACT = """
工程分包合同

发包人：中禾建设集团有限公司
承包人：宏远劳务有限公司

第一条 工程概况
工程名称：某综合楼工程（合成示例）。

第二条 合同价款
本合同价款为固定单价合同，签约合同价 ￥1,286,500.00 元，增值税税率为 9%，价格含税。

第三条 工期
计划开工日期 2026 年 3 月 1 日，工期 180 日历天。

第四条 付款
承包人每月 25 日报送进度款申请，发包人应在收到申请后 28 天内支付已完工程量进度款的 80%。预付款为合同价款的 10%。

第五条 结算
工程竣工验收合格后 60 天内承包人提交竣工结算书，发包人在收到结算书后 90 日内完成结算审核。

第六条 变更与签证
工程变更须经发包人与监理审批后实施，现场签证由发包人、监理、承包人三方共同签认。

第七条 违约责任
承包人工期延误每日历天按合同价款万分之五支付违约金。
""".strip()


@pytest.fixture()
def contract_docx(tmp_path):
    import docx as docx_lib

    d = docx_lib.Document()
    for line in SAMPLE_CONTRACT.splitlines():
        d.add_paragraph(line)
    p = tmp_path / "分包合同.docx"
    d.save(str(p))
    return p


class TestDocxParser:
    def test_paragraphs_extracted(self, contract_docx):
        paras = docx_parser.parse_docx(contract_docx)
        texts = [p["text"] for p in paras]
        assert any("发包人：中禾建设集团有限公司" in t for t in texts)
        assert any("违约" in t for t in texts)


class TestExtract:
    def test_facts_with_quotes(self, contract_docx):
        paras = docx_parser.parse_docx(contract_docx)
        facts = extract.extract_facts(paras)
        by_key = {}
        for f in facts:
            by_key.setdefault(f["fact_key"], []).append(f)
        # 每条事实必须带原文
        for f in facts:
            assert f["quote_text"], "fact must carry quote"
            assert f["location"], "fact must carry location"

        assert any(f["fact_value"] and "中禾建设" in f["fact_value"] for f in by_key.get("employer_party", []))
        amounts = [f["fact_value"] for f in by_key.get("contract_amount", []) if f["fact_value"]]
        assert any("1286500" in a for a in amounts), amounts
        days = [f["fact_value"] for f in by_key.get("payment_clause", []) if f["fact_value"]]
        assert any("28" in d for d in days)
        assert "pricing_method" in by_key
        assert any("固定单价" in (f["fact_value"] or "") for f in by_key["pricing_method"])
        assert "tax_clause" in by_key
        assert "breach_clause" in by_key

    def test_missing_section_not_fabricated(self, tmp_path):
        """没有索赔条款 → 不得编造 claim_clause。"""
        import docx as docx_lib

        d = docx_lib.Document()
        d.add_paragraph("发包人：甲公司")
        d.add_paragraph("承包人：乙公司")
        p = tmp_path / "mini.docx"
        d.save(str(p))
        facts = extract.extract_facts(docx_parser.parse_docx(p))
        keys = {f["fact_key"] for f in facts}
        assert "claim_clause" not in keys
        assert "breach_clause" not in keys
        assert "employer_party" in keys and "contractor_party" in keys


class TestImportAndRisk:
    @pytest.fixture()
    def db(self, tmp_path):
        db_path = tmp_path / "project.db"
        migrations.migrate(db_path, tmp_path / "backups")
        conn = migrations.connect(db_path)
        with conn:
            pid = conn.execute(
                "INSERT INTO projects(name, schema_version, workspace_path, created_at) VALUES ('t',1,'/t','2026')"
            ).lastrowid
        yield conn, pid, tmp_path
        conn.close()

    def test_import_contract_full(self, db, tmp_path):
        conn, pid, pdir = db
        src = tmp_path / "分包合同.docx"
        import docx as docx_lib

        d = docx_lib.Document()
        for line in SAMPLE_CONTRACT.splitlines():
            d.add_paragraph(line)
        d.save(str(src))
        doc_id = extract.import_contract(conn, pid, pdir, src, doc_type="subcontract")
        facts = conn.execute("SELECT fact_key, fact_value, quote_text, evidence_id FROM contract_facts WHERE doc_id=?", (doc_id,)).fetchall()
        assert facts
        assert all(f["quote_text"] and f["evidence_id"] for f in facts)
        # 全量条款 → 无高风险缺失（付款/结算条款都在）
        risks = extract.contract_risks(conn, pid)
        assert not any(r["severity"] == "high" for r in risks), risks

    def test_risk_for_thin_contract(self, db, tmp_path):
        conn, pid, pdir = db
        src = tmp_path / "简短协议.docx"
        import docx as docx_lib

        d = docx_lib.Document()
        d.add_paragraph("发包人：甲公司")
        d.add_paragraph("承包人：乙公司")
        d.save(str(src))
        extract.import_contract(conn, pid, pdir, src)
        risks = extract.contract_risks(conn, pid)
        keys = {r["fact_key"] for r in risks}
        assert {"payment_clause", "settlement_clause", "duration", "contract_amount"} <= keys
        n = extract.persist_risks(conn, pid, risks)
        assert n == len(risks)
        active = run_contract.get_current_contract(conn, pid)
        assert active is not None
        current_scope, scope_params = run_contract.current_scope(conn, pid, "a")
        current_risks = conn.execute(
            f"SELECT COUNT(*) FROM anomalies a WHERE a.project_id=? AND a.rule_id='contract_risk' AND {current_scope}",
            (pid, *scope_params),
        ).fetchone()[0]
        assert current_risks == n
        assert conn.execute(
            "SELECT COUNT(*) FROM evidence WHERE project_id=? AND kind='contract_risk' "
            "AND run_id=?", (pid, active.run_id)
        ).fetchone()[0] == n
        ev = conn.execute("SELECT COUNT(*) c FROM evidence WHERE project_id=?", (pid,)).fetchone()["c"]
        assert ev >= n

        first = conn.execute(
            "SELECT id, evidence_id FROM anomalies WHERE project_id=? AND rule_id='contract_risk' "
            "ORDER BY id LIMIT 1",
            (pid,),
        ).fetchone()
        assert first is not None
        assert extract.persist_risks(conn, pid, []) == 0
        historical = conn.execute(
            "SELECT status, lifecycle_status FROM anomalies WHERE id=?", (first["id"],)
        ).fetchone()
        assert tuple(historical) == ("stale", "historical")
        assert conn.execute(
            "SELECT scope FROM evidence WHERE id=?", (first["evidence_id"],)
        ).fetchone()["scope"] == "historical"
        assert conn.execute(
            "SELECT COUNT(*) FROM finding_status_events WHERE anomaly_id=? AND after_status='historical'",
            (first["id"],),
        ).fetchone()[0] == 1
        assert extract.persist_risks(conn, pid, risks) == n
        current = conn.execute(
            "SELECT repeat_history_json FROM anomalies WHERE project_id=? AND rule_id='contract_risk' "
            "AND lifecycle_status='new' ORDER BY id DESC LIMIT 1",
            (pid,),
        ).fetchone()
        assert current is not None and json.loads(current["repeat_history_json"])

    def test_txt_contract(self, db, tmp_path):
        conn, pid, pdir = db
        src = tmp_path / "会议纪要.txt"
        src.write_text("发包人：甲公司\n支付期限：验收后 30 天内付款\n", encoding="utf-8")
        doc_id = extract.import_contract(conn, pid, pdir, src)
        keys = {r["fact_key"] for r in conn.execute("SELECT fact_key FROM contract_facts WHERE doc_id=?", (doc_id,))}
        assert "payment_clause" in keys

    def test_anomaly_rerun_does_not_hide_current_contract_risks(self, db, tmp_path):
        """通用异常快照重跑不能把合同风险误历史化。"""
        from jiadun.core.anomalies import engine as anomaly_engine

        conn, pid, pdir = db
        src = tmp_path / "简短协议-异常边界.docx"
        import docx as docx_lib

        document = docx_lib.Document()
        document.add_paragraph("发包人：甲公司")
        document.add_paragraph("承包人：乙公司")
        document.save(str(src))
        extract.import_contract(conn, pid, pdir, src)
        risks = extract.contract_risks(conn, pid)
        assert extract.persist_risks(conn, pid, risks) == len(risks)
        before = conn.execute(
            "SELECT COUNT(*) FROM anomalies WHERE project_id=? AND rule_id='contract_risk' "
            "AND lifecycle_status<>'historical'",
            (pid,),
        ).fetchone()[0]
        assert before == len(risks)

        anomaly_engine.run_anomalies(conn, pid, rules=[])

        current_scope, scope_params = run_contract.current_scope(conn, pid, "a")
        after = conn.execute(
            f"SELECT COUNT(*) FROM anomalies a WHERE a.project_id=? "
            f"AND a.rule_id='contract_risk' AND {current_scope}",
            (pid, *scope_params),
        ).fetchone()[0]
        assert after == before
