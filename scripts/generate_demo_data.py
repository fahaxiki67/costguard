"""匿名演示数据生成器（examples/demo）。

生成完全合成的演示结算文件与合同摘录，用于"三分钟上手"与安装包内置演示。
纪律（AGENTS.md）：
- 全部内容为程序合成，不含任何真实公司/项目/人员/金额/业务结构；
- 固定随机种子 + zip 时间戳归一化 → 任意机器重复运行字节一致（--check 验证）；
- 文件与 Sheet 名避开"汇总/核销/台账/summary/reconciliation/ledger"语义门控词，
  唯一例外是"人材机汇总"演示 Sheet——它本身就是角色门控的演示样例。

用法：
  uv run python scripts/generate_demo_data.py            # 生成到 examples/demo
  uv run python scripts/generate_demo_data.py --check    # 重新生成并校验与现有文件字节一致
"""
from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
import tempfile
import zipfile
from datetime import UTC, datetime
from pathlib import Path

import openpyxl
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUT = REPO_ROOT / "examples" / "demo"
sys.path.insert(0, str(REPO_ROOT / "src"))
from jiadun import branding  # noqa: E402

SEED = 20260830
FIXED_TIME = (2026, 1, 1, 0, 0, 0)
FIXED_DATETIME = datetime(2026, 1, 1, 0, 0, 0, tzinfo=UTC)
DEMO_CREATOR = f"{branding.PRODUCT_DISPLAY_NAME}（{branding.PRODUCT_NAME}）合成演示数据生成器"

# 表格样式（演示文件要求：清晰表头、边框、合理列宽）
_HEADER_FONT = Font(bold=True)
_TITLE_FONT = Font(bold=True, size=13)
_HEADER_FILL = PatternFill("solid", fgColor="DDE6F0")
_THIN = Side(style="thin", color="9AA5B1")
_BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)
_NUM_FMT = "#,##0.00"
_COL_WIDTHS = {"A": 15, "B": 22, "C": 26, "D": 9, "E": 12, "F": 12, "G": 14, "H": 8}

# 演示文件名（避开文档级语义门控词；"人材机汇总"仅作为 Sheet 名演示角色门控）
F_UP = "演示-对上结算-第1至3期.xlsx"
F_DOWN = "演示-对下结算-第1至3期.xlsx"
F_DOWN_APPEND = "演示-对下结算-附表.xlsx"
F_CONTRACT = "演示-合同摘录-合成.docx"


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _normalize_zip(path: Path) -> None:
    """把 zip（xlsx/docx）内所有条目的时间戳固定为 FIXED_TIME，保证字节可重复。

    openpyxl 保存时会用当前墙钟时间覆盖 docProps/core.xml 的 dcterms:modified
    （显式赋值也会被覆盖），因此这里对 core.xml 内容做固定替换。
    """
    core_ts = re.compile(
        rb'(<dcterms:(?:created|modified)\b[^>]*>)[^<]*(</dcterms:(?:created|modified)>)')
    fixed_ts = b"2026-01-01T00:00:00Z"
    src = zipfile.ZipFile(path, "r")
    entries = [(info.filename, src.read(info.filename)) for info in src.infolist()]
    src.close()
    tmp = path.with_suffix(path.suffix + ".tmp")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as dst:
        for name, data in entries:
            if name == "docProps/core.xml":
                data = core_ts.sub(rb"\g<1>" + fixed_ts + rb"\g<2>", data)
            info = zipfile.ZipInfo(name, date_time=FIXED_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            info.create_system = 3  # 固定为 Unix，否则 Windows 生成字节与 macOS 不一致
            dst.writestr(info, data)
    # Windows：杀毒软件可能短暂锁定刚写完的文件，os.replace 会 WinError 5，退避重试
    import time

    for attempt in range(6):
        try:
            tmp.replace(path)
            break
        except PermissionError:
            if attempt == 5:
                raise
            time.sleep(0.2 * (attempt + 1))


def _style_table(ws: openpyxl.Worksheet, header_row: int, last_row: int,
                 last_col: int, num_cols: tuple[int, ...]) -> None:
    """表头加粗底色、数据区细边框、数字列右对齐千分位、设置列宽。"""
    for col in range(1, last_col + 1):
        c = ws.cell(row=header_row, column=col)
        c.font = _HEADER_FONT
        c.fill = _HEADER_FILL
        c.border = _BORDER
        c.alignment = Alignment(horizontal="center", vertical="center")
    for row in range(header_row + 1, last_row + 1):
        for col in range(1, last_col + 1):
            c = ws.cell(row=row, column=col)
            c.border = _BORDER
            if col in num_cols:
                c.alignment = Alignment(horizontal="right")
                c.number_format = _NUM_FMT
    for letter, width in _COL_WIDTHS.items():
        if ws.column_dimensions[letter].width is None or True:
            ws.column_dimensions[letter].width = width


def _put_row(ws: openpyxl.Worksheet, row: int, values: list) -> None:
    """values 中 None 表示留空（缺失）；字符串原样写入（可制造文本数字/公式）。"""
    for col, v in enumerate(values, start=1):
        if v is not None:
            ws.cell(row=row, column=col, value=v)


def _title_block(ws: openpyxl.Worksheet, title: str, subtitle: str, end_row: int,
                 merge: bool) -> int:
    """标题区。merge=True 时合并 A1:G1（合并单元格演示）。返回表头行号。"""
    ws.cell(row=1, column=1, value=title).font = _TITLE_FONT
    if merge:
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=7)
    ws.cell(row=2, column=1, value="编制单位：合成演示数据生成器（非真实单位）")
    ws.cell(row=3, column=1, value=subtitle)
    if merge and end_row >= 3:
        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=7)
    return end_row + 1


HEADERS = ["清单编码", "清单名称", "项目特征", "计量单位", "工程量", "综合单价", "合价", "税率"]


def _write_items_sheet(ws: openpyxl.Worksheet, *, title_rows: int, merge_title: bool,
                       subtitle: str, rows: list[list], summary_rows: list[list],
                       blank_after: list[int]) -> None:
    """写一个结算清单 Sheet。rows 为数据行（8 列），summary_rows 为 小计/合计 行。"""
    header_row = _title_block(
        ws, "示例花园小区二期工程（合成演示数据）", subtitle, title_rows, merge_title)
    for col, h in enumerate(HEADERS, start=1):
        ws.cell(row=header_row, column=col, value=h)
    r = header_row
    for i, values in enumerate(rows):
        r += 1
        if i in blank_after:
            r += 1  # 空白行（解析应跳过）
        _put_row(ws, r, values)
    for values in summary_rows:
        r += 1
        _put_row(ws, r, values)
    num_cols = (5, 6, 7, 8) if header_row else (5, 6, 7, 8)
    _style_table(ws, header_row, r, 8, num_cols)
    for letter, width in _COL_WIDTHS.items():
        ws.column_dimensions[letter].width = width


def build_upward(path: Path) -> None:
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    # ---- 第1期：表头在第 4 行 + 合并标题；覆盖 文本数字/千分位/真零/缺失/负数/不符/舍入/公式 ----
    ws = wb.create_sheet("第1期")
    rows = [
        ["010101001001", "平整场地", "三类土，人工配合", "m2", 1250.50, 8.50, 10629.25, "9%"],
        ["010101003001", "挖沟槽土方", "深度2m以内，三类土", "m3", 860.00, 35.20, 30272.00, "9%"],
        ["010401001001", "砖基础", "MU10机制砖，M7.5水泥砂浆", "m3", 152.30, 420.00, 63966.00, "9%"],
        ["010501001001", "C25混凝土垫层", "商品混凝土C25，泵送", "m3", 48.20, 465.00, 22413.00, "9%"],
        ["010515001001", "现浇构件钢筋", "HRB400，直径10mm以内", "t", 25.80, 4850.00, 125130.00, "13%"],
        ["011001001001", "水泥砂浆楼地面", "1:3水泥砂浆，20mm厚", "m2", "680.00", 45.80, 31144.00, "9%"],
        ["010416001001", "钢筋接头电渣压力焊", "直径14mm以内", "个", 0, 3.50, 0.00, "9%"],
        ["010901001001", "屋面SBS防水卷材", "4mm厚，热熔施工", "m2", 320.00, None, None, "9%"],
        ["011702002001", "脚手架搭拆费", "钢管双排脚手架，冲减", "m2", -150.00, 12.00, -1800.00, "9%"],
        ["010402001001", "蒸压加气混凝土砌块墙", "B06级，专用砂浆", "m3", 96.40, 320.00, 31848.00, "9%"],
        ["010503001001", "现浇混凝土构造柱", "C25，截面240x240", "m3", 36.80, 520.00, 19136.01, "9%"],
        ["010515002001", "预埋铁件", "综合铁件", "kg", 200.00, 6.80, "=E14*F14", "9%"],
        ["011703001001", "安全文明施工费", "费率计取", "项", 1.00, 25000.00, "=E15*F15+", "9%"],
    ]
    subtotal = sum(v for v in [10629.25, 30272.00, 63966.00, 22413.00, 125130.00, 31144.00,
                               0.00, None, -1800.00, 31848.00, 19136.01, None, None]
                   if v is not None)
    _write_items_sheet(
        ws, title_rows=3, merge_title=True, subtitle="对上结算 第1期（合成演示）",
        rows=rows,
        summary_rows=[[None, "小计", None, None, None, None, round(subtotal, 2), None],
                      [None, "合计", None, None, None, None, round(subtotal, 2), None]],
        blank_after=[],
    )

    # ---- 第2期：表头在第 1 行；同码异名/单价变化/税率变化/单位变化/重复行/近似不同项 ----
    ws = wb.create_sheet("第2期")
    rows = [
        ["010101001001", "平整场地", "三类土", "m2", 400.00, 8.50, 3400.00, "9%"],
        ["010101003001", "挖沟槽土方（机械）", "深度2m以内", "m3", 300.00, 46.00, 13800.00, "9%"],
        ["010401001001", "砖基础", "MU10机制砖，M7.5水泥砂浆", "m3", 60.00, 425.00, 25500.00, "13%"],
        ["010501001001", "C25混凝土垫层", "商品混凝土C25，泵送", "m2", 120.00, 78.00, 9360.00, "9%"],
        ["010515001001", "现浇构件钢筋", "HRB400，直径10mm以内", "t", 8.00, 4850.00, 38800.00, "13%"],
        ["010515001001", "现浇构件钢筋", "HRB400，直径10mm以内", "t", 8.00, 4850.00, 38800.00, "13%"],
        ["010501002001", "C30混凝土垫层", "商品混凝土C30，泵送", "m3", 55.00, 480.00, 26400.00, "9%"],
        ["011001002001", "保温楼地面", "挤塑聚苯板50mm厚", "m2", 210.00, 92.00, 19320.00, "9%"],
    ]
    subtotal = 3400.00 + 13800.00 + 25500.00 + 9360.00 + 38800.00 + 38800.00 + 26400.00 + 19320.00
    _write_items_sheet(
        ws, title_rows=0, merge_title=False, subtitle="对上结算 第2期（合成演示）",
        rows=rows,
        summary_rows=[[None, "小计", None, None, None, None, round(subtotal, 2), None],
                      [None, "合计", None, None, None, None, round(subtotal, 2), None]],
        blank_after=[],
    )

    # ---- 第3期：漏项（钢筋/挖土/防水缺失）、工程量突增、缺失单价金额、空白行 ----
    ws = wb.create_sheet("第3期")
    rows = [
        ["010101001001", "平整场地", "三类土", "m2", 5000.00, 8.50, 42500.00, "9%"],
        ["010401001001", "砖基础", "MU10机制砖", "m3", 20.00, 425.00, 8500.00, "13%"],
        ["010501001001", "C25混凝土垫层", "商品混凝土C25", "m2", 80.00, 78.00, 6240.00, "9%"],
        ["012999001001", "其他零星工作", "综合", "项", 1.00, None, None, None],
    ]
    _write_items_sheet(
        ws, title_rows=1, merge_title=True, subtitle="对上结算 第3期（合成演示）",
        rows=rows,
        summary_rows=[[None, "小计", None, None, None, None, 57240.00, None],
                      [None, "合计", None, None, None, None, 57240.00, None]],
        blank_after=[2],
    )
    wb.properties.creator = DEMO_CREATOR
    wb.properties.lastModifiedBy = DEMO_CREATOR
    wb.properties.created = FIXED_DATETIME
    wb.properties.modified = FIXED_DATETIME
    wb.save(path)
    _normalize_zip(path)


def build_downward(path: Path) -> None:
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    # ---- 第1期（对下）：与对上同期号 → 方向隔离演示；缺失数量行 ----
    ws = wb.create_sheet("第1期")
    rows = [
        ["010101001001", "平整场地", "三类土", "m2", 1250.50, 6.80, 8503.40, "9%"],
        ["010401001001", "砖基础", "MU10机制砖", "m3", 152.30, 385.00, 58635.50, "9%"],
        ["010501001001", "C25混凝土垫层", "商品混凝土C25", "m3", 48.20, 430.00, 20726.00, "9%"],
        ["D1-001", "劳务队机械土方", "综合机械配合", "m3", 860.00, 18.00, 15480.00, "9%"],
        [None, "劳务队土方用工", "综合人工配合", "m3", 200.00, 18.00, 3600.00, "9%"],
        ["010416001001", "钢筋接头电渣压力焊", "直径14mm以内", "个", None, 3.20, None, "9%"],
    ]
    subtotal = 8503.40 + 58635.50 + 20726.00 + 15480.00 + 3600.00
    _write_items_sheet(
        ws, title_rows=1, merge_title=True, subtitle="对下结算 第1期（合成演示）",
        rows=rows,
        summary_rows=[[None, "小计", None, None, None, None, round(subtotal, 2), None],
                      [None, "合计", None, None, None, None, round(subtotal, 2), None]],
        blank_after=[],
    )

    # ---- 第2期（对下）：无编码近名行 → 独立高概率组；缺名称行 → 待补资料档演示 ----
    ws = wb.create_sheet("第2期")
    rows = [
        ["010101001001", "平整场地", "三类土", "m2", 400.00, 6.80, 2720.00, "9%"],
        ["010401001001", "砖基础", "MU10机制砖", "m3", 60.00, 388.00, 23280.00, "9%"],
        [None, "C25砼垫层", "商品混凝土", "m3", 40.00, 430.00, 17200.00, "9%"],
        [None, None, None, "m2", 30.00, 45.00, 1350.00, "9%"],
    ]
    subtotal = 2720.00 + 23280.00 + 17200.00 + 1350.00
    _write_items_sheet(
        ws, title_rows=0, merge_title=False, subtitle="对下结算 第2期（合成演示）",
        rows=rows,
        summary_rows=[[None, "小计", None, None, None, None, round(subtotal, 2), None],
                      [None, "合计", None, None, None, None, round(subtotal, 2), None]],
        blank_after=[],
    )

    # ---- 第3期（对下）----
    ws = wb.create_sheet("第3期")
    rows = [
        ["010101001001", "平整场地", "三类土", "m2", 5000.00, 6.80, 34000.00, "9%"],
        ["010401001001", "砖基础", "MU10机制砖", "m3", 20.00, 388.00, 7760.00, "9%"],
    ]
    subtotal = 34000.00 + 7760.00
    _write_items_sheet(
        ws, title_rows=1, merge_title=True, subtitle="对下结算 第3期（合成演示）",
        rows=rows,
        summary_rows=[[None, "小计", None, None, None, None, round(subtotal, 2), None],
                      [None, "合计", None, None, None, None, round(subtotal, 2), None]],
        blank_after=[],
    )
    wb.properties.creator = DEMO_CREATOR
    wb.properties.lastModifiedBy = DEMO_CREATOR
    wb.properties.created = FIXED_DATETIME
    wb.properties.modified = FIXED_DATETIME
    wb.save(path)
    _normalize_zip(path)


def build_downward_append(path: Path) -> None:
    """附表：正常补充期次 Sheet + "人材机汇总"角色门控演示 Sheet。"""
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    ws = wb.create_sheet("第1期补充")
    rows = [
        ["011001001001", "水泥砂浆楼地面", "1:3水泥砂浆，20mm厚", "m2", 680.00, 38.00, 25840.00, "9%"],
        ["011702002001", "脚手架搭拆费", "钢管双排脚手架", "m2", 150.00, 9.00, 1350.00, "9%"],
    ]
    _write_items_sheet(
        ws, title_rows=1, merge_title=True, subtitle="对下结算 第1期补充（合成演示）",
        rows=rows,
        summary_rows=[[None, "小计", None, None, None, None, 27190.00, None],
                      [None, "合计", None, None, None, None, 27190.00, None]],
        blank_after=[],
    )

    # 人材机汇总型表：Sheet 名含"汇总"→ 语义门控，导入后进入"待人工角色确认"。
    # 这是刻意的演示：展示价盾不猜测、不给权限就写入正式清单。
    ws = wb.create_sheet("人材机汇总")
    ws.cell(row=1, column=1, value="人材机汇总表（合成演示，导入后需人工确认角色）").font = _TITLE_FONT
    headers = ["序号", "名称", "单位", "数量", "单价", "金额"]
    for col, h in enumerate(headers, start=1):
        ws.cell(row=2, column=col, value=h)
    data = [
        [1, "综合工日", "工日", 1200.00, 260.00, 312000.00],
        [2, "商品混凝土C25", "m3", 143.20, 440.00, 63008.00],
        [3, "HRB400钢筋", "t", 33.80, 4780.00, 161564.00],
    ]
    for i, row in enumerate(data, start=3):
        _put_row(ws, i, row)
    _style_table(ws, 2, 5, 6, (4, 5, 6))
    for letter, width in {"A": 8, "B": 20, "C": 9, "D": 12, "E": 12, "F": 14}.items():
        ws.column_dimensions[letter].width = width

    wb.properties.creator = DEMO_CREATOR
    wb.properties.lastModifiedBy = DEMO_CREATOR
    wb.properties.created = FIXED_DATETIME
    wb.properties.modified = FIXED_DATETIME
    wb.save(path)
    _normalize_zip(path)


CONTRACT_PARAGRAPHS = [
    "第一部分 合同价格形式",
    "1.1 本合同采用单价合同形式。综合单价包含人工费、材料费、机械费、管理费、利润和风险费用。",
    "1.2 综合单价在合同约定的风险范围内固定不变，超出范围时按合同专用条款第 3.2 款调整。",
    "第二部分 进度款与支付",
    "2.1 进度款按月支付，支付比例为当月经确认已完成工程量价款的 85%。",
    "2.2 发包人应在收到承包人付款申请后 14 天内完成审核并支付。",
    "第三部分 变更与估价",
    "3.1 变更估价参照已标价清单中相同或类似项目单价执行。",
    "3.2 已标价清单中无适用项目时，由双方按照成本加合理利润原则协商确定，并形成书面记录。",
    "第四部分 竣工结算",
    "4.1 承包人应在工程竣工验收合格后 28 天内提交竣工结算申请单及完整结算资料。",
    "4.2 发包人应在收到结算资料后 60 天内完成审核；逾期未提出异议的，视为认可承包人提交的结算文件。",
    "第五部分 质量保证金",
    "5.1 缺陷责任期为 24 个月，自工程实际竣工验收合格之日起计算。",
    "5.2 质量保证金按工程价款结算总额的 3% 预留，缺陷责任期满后 14 天内无息返还。",
    "第六部分 争议解决（合成条款）",
    "6.1 因本合同发生的争议，双方协商解决；协商不成的，提交示例市仲裁委员会仲裁（合成表述，非真实机构）。",
]


def build_contract(path: Path) -> None:
    import docx

    doc = docx.Document()
    doc.add_heading("示例花园小区二期工程 施工合同（摘录·合成演示）", level=0)
    doc.add_paragraph(
        "本文件为程序生成的合成演示文档，全部条款内容均为虚构，不代表任何真实合同或业务结论。")
    for text in CONTRACT_PARAGRAPHS:
        if text.startswith("第") and "部分" in text[:6]:
            doc.add_heading(text, level=1)
        else:
            doc.add_paragraph(text)
    props = doc.core_properties
    props.author = DEMO_CREATOR
    props.last_modified_by = DEMO_CREATOR
    props.created = FIXED_DATETIME
    props.modified = FIXED_DATETIME
    doc.save(path)
    _normalize_zip(path)


# ---------------------------------------------------------------------------
# 期望值（manifest）：预期值以实际导入管线行为为准，并由 tests/unit/test_demo_data.py 锁定
# ---------------------------------------------------------------------------
EXPECTATIONS = {
    F_UP: {
        "data_type": "xlsx", "direction": "upward", "periods": [1, 2, 3],
        "expected_parsed_rows": {"第1期": 13, "第2期": 8, "第3期": 4},
        "expected_anomalies": [
            "qty_price_amount_mismatch", "rounding_difference", "text_number_in_value_col",
            "formula_no_cache", "unparsed_number", "negative_quantity", "duplicate_item",
            "price_changed", "price_abnormal_change", "tax_rate_changed", "unit_changed",
            "qty_sudden_change", "missing_key_fields", "summary_mismatch",
            "summary_missing_data", "large_round_amount", "same_code_diff_name",
        ],
        "expected_matching": "同码同名=规则完全匹配（待人工确认）；同码异名/近似名进入人工复核",
        "coverage": ["多期", "多Sheet", "不同表头位置", "合并单元格", "文本数字", "千分位",
                     "真实零值", "缺失不填零", "负数调整", "数量单价合价不符", "一分钱舍入差异",
                     "同码名称轻微不同", "名称相似实际不同", "单位不一致", "单价变化", "税率变化",
                     "重复项", "漏项", "公式单元格", "公式缺少缓存值", "小计合计行"],
        "known_limitations": [
            "第1期两条公式行无缓存值，导入后合价显示待补资料并触发 formula_no_cache",
            "合并单元格仅用于标题区，数据区不使用合并单元格",
        ],
    },
    F_DOWN: {
        "data_type": "xlsx", "direction": "downward", "periods": [1, 2, 3],
        "expected_parsed_rows": {"第1期": 6, "第2期": 4, "第3期": 2},
        "expected_anomalies": ["missing_key_fields", "price_changed", "orphan_numeric_row"],
        "expected_matching": "无编码行按名称独立成组（高概率匹配）；缺名称行为待补资料档",
        "coverage": ["对下方向", "与对上同期号（方向隔离演示）", "匹配置信度多档",
                     "自定义编码（D1-001）", "需要人工复核"],
        "known_limitations": ["对下单价低于对上为演示设定，无调差含义"],
    },
    F_DOWN_APPEND: {
        "data_type": "xlsx", "direction": "downward", "periods": [1],
        "expected_parsed_rows": {"第1期补充": 2},
        "expected_anomalies": [],
        "expected_matching": "补充 Sheet 并入对下第1期后参与累计",
        "coverage": ["多文件同期次合并", "角色门控演示（人材机汇总 Sheet 待人工确认）"],
        "known_limitations": [
            "Sheet「人材机汇总」被语义门控拦截属预期行为：演示程序不猜测角色，"
            "需人工确认后才写入结算模型",
        ],
    },
    F_CONTRACT: {
        "data_type": "docx", "direction": "contract", "periods": [],
        "expected_parsed_rows": {"paragraphs": len(CONTRACT_PARAGRAPHS)},
        "expected_anomalies": [
            "未识别工期约定（medium 风险提示）",
            "未识别税务条款（medium 风险提示）",
            "条款提取结果须逐条人工核对原文引用",
        ],
        "expected_matching": "不参与清单匹配；风险提示为程序识别结果，不是业务结论",
        "coverage": ["合同演示文本", "完全合成的原文引用", "合同风险提示演示"],
        "known_limitations": ["条款提取为初步能力，自动结果不等于业务结论"],
    },
}

DISCLAIMER = ("本目录全部内容为 scripts/generate_demo_data.py 程序生成的合成演示数据，"
              "不包含任何真实公司、项目、人员、金额或业务结构，不代表真实业务结论，"
              "仅用于功能演示与上手练习。")

COVERAGE_MATRIX = {
    "对上结算与对下结算": F_UP + " + " + F_DOWN,
    "同项目同期号不同方向（方向隔离）": "两个文件均含 第1/2/3期",
    "多期 Excel": "两个结算文件各含 3 期",
    "多 Sheet": "附表文件含 2 个 Sheet",
    "不同表头位置": "表头分别位于第 1/2/4 行",
    "合并单元格": "标题行合并（数据区不用合并，避免解析歧义）",
    "文本数字": "对上第1期 水泥砂浆楼地面 工程量为文本 680.00",
    "千分位金额": "数字列统一千分位格式，含 125,130.00 大额行",
    "真实零值": "对上第1期 钢筋接头行 工程量 0、合价 0",
    "缺失数量/单价/金额": "对上第1期缺单价金额行；第3期零星工作行；对下第1期缺数量行",
    "缺失不填零": "缺失单元格导入后为待补资料，程序不自动补 0",
    "负数调整项": "对上第1期 脚手架搭拆费冲减行",
    "数量x单价与合价不一致": "对上第1期 砌块墙行（差 1000.00）",
    "一分钱舍入差异": "对上第1期 构造柱行（差 0.01）",
    "同编码名称轻微不同": "挖沟槽土方 → 挖沟槽土方（机械）",
    "名称相似实际不同": "C25混凝土垫层 与 C30混凝土垫层",
    "单位不一致": "C25混凝土垫层 第1期 m3 → 第2期 m2",
    "单价变化": "砖基础 420→425；挖沟槽土方 35.20→46.00（超阈值升级）",
    "税率变化": "砖基础 9% → 13%",
    "重复项": "对上第2期 现浇构件钢筋 完全重复两行",
    "漏项": "对上第3期缺 现浇构件钢筋/挖沟槽土方/屋面防水",
    "公式单元格": "对上第1期 预埋铁件行",
    "公式错误或缺少缓存值": "对上第1期 安全文明施工费行（无缓存）",
    "小计合计行不重复计入": "各 Sheet 末尾 小计/合计 行，导入打小计标记",
    "合同演示文本": F_CONTRACT + "（条款全部虚构）",
    "匹配置信度多档": "规则完全匹配（待人工确认）/疑似匹配/待补资料",
    "需要人工复核的项目": "人材机汇总 Sheet 角色门控 + 各缺失行",
    "可导出 Excel 和 Word": "导入演示数据后可从成果导出页导出审核底稿与摘要",
}


def generate(out_dir: Path) -> dict:
    """生成全部演示文件并返回 manifest dict。"""
    out_dir.mkdir(parents=True, exist_ok=True)
    build_upward(out_dir / F_UP)
    build_downward(out_dir / F_DOWN)
    build_downward_append(out_dir / F_DOWN_APPEND)
    build_contract(out_dir / F_CONTRACT)

    files = []
    for name in (F_UP, F_DOWN, F_DOWN_APPEND, F_CONTRACT):
        exp = EXPECTATIONS[name]
        files.append({
            "file_name": name,
            "sha256": _sha256(out_dir / name),
            **exp,
        })
    manifest = {
        "schema_version": 1,
        "generated_by": "scripts/generate_demo_data.py",
        "random_seed": SEED,
        "generated_at_fixed": "2026-01-01T00:00:00+00:00（时间戳固定，实际内容可重复生成）",
        "disclaimer": DISCLAIMER,
        "project_profile": {
            "project_name": "示例花园小区二期工程（虚构）",
            "parties": "演示发包人 / 演示承包人 / 演示劳务分包（均为虚构）",
        },
        "files": files,
        "coverage_matrix": COVERAGE_MATRIX,
    }
    (out_dir / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    sums = "\n".join(f"{_sha256(out_dir / name)}  {name}" for name in
                     (F_UP, F_DOWN, F_DOWN_APPEND, F_CONTRACT)) + "\n"
    (out_dir / "SHA256SUMS").write_text(sums, encoding="utf-8")
    return manifest


def main() -> int:
    parser = argparse.ArgumentParser(description="生成匿名演示数据（examples/demo）")
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT, help="输出目录")
    parser.add_argument("--check", action="store_true",
                        help="重新生成到临时目录，校验与现有文件字节一致（确定性检查）")
    args = parser.parse_args()

    if args.check:
        target = args.out
        if not target.exists():
            print(f"FAIL: 目标目录不存在：{target}", file=sys.stderr)
            return 1
        with tempfile.TemporaryDirectory() as td:
            tmp = Path(td) / "demo"
            generate(tmp)
            bad = []
            for name in (F_UP, F_DOWN, F_DOWN_APPEND, F_CONTRACT,
                         "manifest.json", "SHA256SUMS"):
                a, b = target / name, tmp / name
                if not a.exists() or _sha256(a) != _sha256(b):
                    bad.append(name)
            if bad:
                print("FAIL: 以下文件与重新生成结果不一致（演示数据必须确定性）：", file=sys.stderr)
                for n in bad:
                    print(f"  {n}", file=sys.stderr)
                return 1
        print(f"OK: 演示数据确定性校验通过（{target}）")
        return 0

    generate(args.out)
    print(f"已生成演示数据：{args.out}")
    for name in (F_UP, F_DOWN, F_DOWN_APPEND, F_CONTRACT):
        p = args.out / name
        print(f"  {name}  {_sha256(p)[:16]}…  {p.stat().st_size} bytes")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
