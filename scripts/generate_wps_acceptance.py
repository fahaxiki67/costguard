"""生成最小 WPS 人工验收文件（监督指令三）。

产物写入 local_private_data/wps_acceptance/（不入 Git）：
- WPS验收-最小底稿.xlsx：《审核底稿》《汇总校验》两表，数据区细边框、
  水平垂直居中、表头加粗统一底色、长文本自动换行、列宽适配 WPS 阅读；
  保留公式、Decimal 精度、缺失值标记、负数与真零语义
- WPS验收-最小摘要.docx

这些文件由程序合成，仅用于 WPS 打开验收，不含任何真实业务数据。
"""
from __future__ import annotations

from decimal import Decimal
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

HEADER_FILL = PatternFill("solid", fgColor="DDEBF7")
MONEY_FMT = "#,##0.00"
_THIN = Side(style="thin", color="FF808080")
DATA_BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)
CENTER_WRAP = Alignment(horizontal="center", vertical="center", wrap_text=True)
HEADER_FONT = Font(bold=True)

# 列宽（字符数）——按 WPS 阅读习惯
AUDIT_WIDTHS = (15, 26, 7, 11, 11, 14, 14, 14, 18)
CHECK_WIDTHS = (26, 18)


def _style_region(ws, n_rows: int, n_cols: int, widths: tuple[int, ...],
                  center_cols: tuple[int, ...] | None = None) -> None:
    """表头加粗+统一底色；数据区细边框+水平垂直居中+自动换行；列宽设置。

    只改外观属性（font/fill/alignment/border/column width），
    不触碰单元格值、公式与 number_format——业务语义零改动。
    """
    for c, w in enumerate(widths[:n_cols], start=1):
        ws.column_dimensions[get_column_letter(c)].width = w
    for r in range(1, n_rows + 1):
        for c in range(1, n_cols + 1):
            cell = ws.cell(row=r, column=c)
            cell.border = DATA_BORDER
            cell.alignment = CENTER_WRAP
            if r == 1:
                cell.fill = HEADER_FILL
                cell.font = HEADER_FONT


def generate(out_dir: Path) -> tuple[Path, Path]:
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    wb = Workbook()
    wb.remove(wb.active)
    ws = wb.create_sheet("审核底稿")
    ws.append(["清单编码", "清单名称", "单位", "数量", "单价", "合价(公式)", "原表合价", "差异(公式)", "状态"])
    rows = [
        # (编码, 名称, 单位, 数量, 单价, 原表合价, 状态)
        ("010101001001", "平整场地（验收样例）", "m2", "1250.50", "8.50", "10629.25", "正常"),
        ("010501001001", "混凝土垫层-精度检查", "m3", "0.1", "465.00", "46.50", "Decimal 精度"),
        ("010401001001", "真零检查行", "m3", "0", "420.00", "0", "真零（应为 0 不是空）"),
        ("010515001001", "缺失单价行", "t", "25.80", None, "125130.00", "待补资料"),
        ("011001001001", "负数冲减行", "m2", "-50", "45.80", "-2290.00", "负数"),
        ("011001002001", "千分位大额行", "m2", "680.00", "45.80", "31144.00", "金额格式"),
    ]
    for i, (code, name, unit, qty, price, amount, status) in enumerate(rows, start=2):
        r = i
        ws.cell(row=r, column=1, value=code)
        ws.cell(row=r, column=2, value=name)
        ws.cell(row=r, column=3, value=unit)
        ws.cell(row=r, column=4, value=Decimal(qty))  # Decimal 直写（真零保留）
        if price is None:
            ws.cell(row=r, column=5, value="待补资料")  # 缺失标记，不补 0
        else:
            ws.cell(row=r, column=5, value=Decimal(price))
            ws.cell(row=r, column=6, value=f"=D{r}*E{r}")  # 合价公式
        ws.cell(row=r, column=7, value=Decimal(amount))  # 负数原值保留
        if price is not None:
            ws.cell(row=r, column=8, value=f"=ROUND(F{r}-G{r},2)")  # 差异公式
        else:
            ws.cell(row=r, column=8, value="不可比")
        ws.cell(row=r, column=9, value=status)
        for c in (4, 5, 6, 7, 8):
            ws.cell(row=r, column=c).number_format = MONEY_FMT
    _style_region(ws, n_rows=ws.max_row, n_cols=9, widths=AUDIT_WIDTHS)

    ws2 = wb.create_sheet("汇总校验")
    ws2.append(["校验项", "数值(公式)"])
    ws2.cell(row=2, column=1, value="明细原表合价合计")
    ws2.cell(row=2, column=2, value="=SUM(审核底稿!G2:G7)")
    ws2.cell(row=3, column=1, value="公式合价合计")
    ws2.cell(row=3, column=2, value="=SUM(审核底稿!F2:F7)")
    ws2.cell(row=4, column=1, value="差异合计")
    ws2.cell(row=4, column=2, value="=B3-B2")
    for r in (2, 3, 4):
        ws2.cell(row=r, column=2).number_format = MONEY_FMT
    _style_region(ws2, n_rows=ws2.max_row, n_cols=2, widths=CHECK_WIDTHS)

    xlsx_path = out_dir / "WPS验收-最小底稿.xlsx"
    wb.save(xlsx_path)

    import docx as docx_lib

    doc = docx_lib.Document()
    doc.add_heading("CostGuard WPS 验收-最小摘要", level=0)
    doc.add_paragraph("本文件由程序生成，仅用于验证 WPS 打开、中文显示与数字格式，"
                      "不含任何真实业务数据。")
    doc.add_paragraph("验收要点：标题可见、段落中文正常、不乱码。")
    docx_path = out_dir / "WPS验收-最小摘要.docx"
    doc.save(str(docx_path))
    return xlsx_path, docx_path


def count_formulas_and_errors(xlsx_path: Path, recalc_dir: Path | None = None) -> dict:
    """统计公式数量；若提供 LibreOffice 重算副本目录，统计重算错误数。"""
    import openpyxl

    wb = openpyxl.load_workbook(xlsx_path, data_only=False)
    formulas = []
    for name in ("审核底稿", "汇总校验"):
        ws = wb[name]
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.startswith("="):
                    formulas.append(f"{name}!{cell.coordinate}:{cell.value}")
    errors = None
    if recalc_dir is not None:
        recalc = recalc_dir / xlsx_path.name
        wb2 = openpyxl.load_workbook(recalc, data_only=True)
        errors = 0
        for name in ("审核底稿", "汇总校验"):
            ws2 = wb2[name]
            for row in ws2.iter_rows():
                for cell in row:
                    v = cell.value
                    if isinstance(v, str) and v.startswith("#"):
                        errors += 1
    return {"n_formulas": len(formulas), "formulas": formulas, "recalc_errors": errors}


if __name__ == "__main__":
    out = Path.home() / ".zcode/workspace/default/costguard/local_private_data/wps_acceptance"
    x, d = generate(out)
    print(f"xlsx: {x}\ndocx: {d}")
